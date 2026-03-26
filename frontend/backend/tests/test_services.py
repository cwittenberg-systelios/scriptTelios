"""
Tests fuer Service-Schicht: llm, job_queue, docx_fill, embeddings.

Fokus auf bisher unabgedeckte Pfade:
  - llm.py:       Timeout, HTTPStatusError, leere Response
  - job_queue.py: Error-Pfad, Cleanup, to_dict-Felder
  - docx_fill.py: find_placeholders, Platzhalter-Ersetzung, Append-Pfad
  - embeddings.py: get_embedding Fehlerfaelle, retrieve_style_examples Fallback
"""
import asyncio
import uuid
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import httpx
import pytest

# ══════════════════════════════════════════════════════════════════
# LLM SERVICE
# ══════════════════════════════════════════════════════════════════

class TestLLMService:

    @pytest.mark.asyncio
    async def test_generate_text_erfolgreich(self):
        """generate_text gibt text, model_used und duration_s zurueck."""
        mock_response = MagicMock()
        mock_response.json.return_value = {
            "response": "Generierter Text.",
            "eval_count": 42,
        }
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.llm import generate_text
            result = await generate_text("System", "User")

        assert result["text"] == "Generierter Text."
        assert "ollama/" in result["model_used"]
        assert result["token_count"] == 42
        assert "duration_s" in result

    @pytest.mark.asyncio
    async def test_generate_text_connect_error_wirft_runtime(self):
        """ConnectError wird als RuntimeError mit sprechender Meldung weitergegeben."""
        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                side_effect=httpx.ConnectError("refused")
            )
            from app.services.llm import generate_text
            with pytest.raises(RuntimeError, match="Ollama nicht erreichbar"):
                await generate_text("System", "User")

    @pytest.mark.asyncio
    async def test_generate_text_http_status_error(self):
        """HTTPStatusError (z.B. 500) wird als RuntimeError weitergegeben."""
        mock_response = MagicMock()
        mock_response.status_code = 500
        mock_response.text = "Internal Server Error"

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                side_effect=httpx.HTTPStatusError(
                    "500", request=MagicMock(), response=mock_response
                )
            )
            from app.services.llm import generate_text
            with pytest.raises(RuntimeError, match="Ollama Fehler 500"):
                await generate_text("System", "User")

    @pytest.mark.asyncio
    async def test_generate_text_leere_response(self):
        """Leere Ollama-Response liefert leeren Text ohne Fehler."""
        mock_response = MagicMock()
        mock_response.json.return_value = {"response": ""}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.llm import generate_text
            result = await generate_text("System", "User")

        assert result["text"] == ""

    @pytest.mark.asyncio
    async def test_generate_text_oom_retry_mit_reduziertem_kontext(self):
        """Bei VRAM-OOM: Modell entladen, zweiter Versuch mit 8192 Tokens erfolgreich."""
        oom_response = MagicMock()
        oom_response.status_code = 500
        oom_response.text = "cuda out of memory"

        ok_response = MagicMock()
        ok_response.json.return_value = {"response": "Text mit kleinem Kontext.", "eval_count": 10}
        ok_response.raise_for_status = MagicMock()

        call_count = 0

        async def mock_post(url, json=None, **kwargs):
            nonlocal call_count
            call_count += 1
            if call_count == 1:
                # Erster Aufruf: OOM
                raise httpx.HTTPStatusError("500", request=MagicMock(), response=oom_response)
            if call_count == 2:
                # Unload-Aufruf (keep_alive=0)
                return MagicMock(raise_for_status=MagicMock())
            # Dritter Aufruf: Retry mit reduziertem Kontext → Erfolg
            return ok_response

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(side_effect=mock_post)
            from app.services import llm as llm_module
            import importlib
            importlib.reload(llm_module)
            result = await llm_module.generate_text("System", "User")

        assert result["text"] == "Text mit kleinem Kontext."

    @pytest.mark.asyncio
    async def test_generate_text_oom_retry_schlaegt_auch_fehl(self):
        """Bei OOM auf beiden Versuchen: klare Fehlermeldung mit Hinweis auf .env."""
        oom_response = MagicMock()
        oom_response.status_code = 500
        oom_response.text = "cuda out of memory"

        async def mock_post(url, json=None, **kwargs):
            if json and json.get("keep_alive") == 0:
                return MagicMock(raise_for_status=MagicMock())  # Unload gelingt
            raise httpx.HTTPStatusError("500", request=MagicMock(), response=oom_response)

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(side_effect=mock_post)
            from app.services import llm as llm_module
            import importlib
            importlib.reload(llm_module)
            with pytest.raises(RuntimeError, match="WHISPER_FREE_OLLAMA_VRAM"):
                await llm_module.generate_text("System", "User")

    def test_deduplicate_paragraphs_entfernt_wiederholungen(self):
        """Doppelte Absätze werden aus dem LLM-Output entfernt."""
        from app.services.llm import deduplicate_paragraphs
        text = "Erster Absatz.\n\nZweiter Absatz.\n\nErster Absatz.\n\nDritter Absatz."
        result = deduplicate_paragraphs(text)
        assert result.count("Erster Absatz.") == 1
        assert "Zweiter Absatz." in result
        assert "Dritter Absatz." in result

    def test_deduplicate_paragraphs_case_insensitive(self):
        """Duplikat-Erkennung ist case-insensitiv."""
        from app.services.llm import deduplicate_paragraphs
        text = "Die Klientin kam mit dem Anliegen.\n\ndie klientin kam mit dem anliegen."
        result = deduplicate_paragraphs(text)
        assert result.count("Klientin") + result.count("klientin") == 1

    def test_deduplicate_paragraphs_behaelt_einzigartigen_text(self):
        """Unique Absätze bleiben vollständig erhalten."""
        from app.services.llm import deduplicate_paragraphs
        text = "Abschnitt A.\n\nAbschnitt B.\n\nAbschnitt C."
        assert deduplicate_paragraphs(text) == text

    def test_deduplicate_paragraphs_leerer_input(self):
        """Leerer Input gibt leeren String zurück."""
        from app.services.llm import deduplicate_paragraphs
        assert deduplicate_paragraphs("") == ""

    def test_truncate_style_context_kurzer_text_unveraendert(self):
        """Texte unter dem Limit werden nicht gekürzt."""
        from app.services.llm import truncate_style_context, MAX_STYLE_CONTEXT_CHARS
        short = "Kurzer Stiltext." * 10
        assert len(short) < MAX_STYLE_CONTEXT_CHARS
        assert truncate_style_context(short) == short

    def test_truncate_style_context_kuerzt_an_satzgrenze(self):
        """Langer Text wird an einer Satzgrenze gekürzt, nicht mitten im Satz."""
        from app.services.llm import truncate_style_context, MAX_STYLE_CONTEXT_CHARS
        # Erzeuge Text der deutlich über dem Limit liegt
        long_text = ("Dies ist ein Satz der Informationen enthält. " * 100)
        result = truncate_style_context(long_text)
        assert len(result) <= MAX_STYLE_CONTEXT_CHARS
        # Muss mit Satzzeichen enden
        assert result[-1] in ".!?"

    def test_truncate_style_context_repetitiver_input(self):
        """Repetitive C&P-Stilvorlage (wie im Bug-Report) wird auf Limit gekürzt."""
        from app.services.llm import truncate_style_context, MAX_STYLE_CONTEXT_CHARS
        # Simuliert die eingefügte Stilvorlage die 8x denselben Absatz enthielt
        absatz = "Die Klientin kam mit dem Anliegen, ihre Schwierigkeiten zu bewältigen. " * 5
        repetitiv = absatz * 8
        result = truncate_style_context(repetitiv)
        assert len(result) <= MAX_STYLE_CONTEXT_CHARS

    @pytest.mark.asyncio
    async def test_generate_text_dedupliziert_output(self):
        """generate_text bereinigt automatisch doppelte Absätze im LLM-Output."""
        wiederholter_output = (
            "Erster Abschnitt mit Inhalt.\n\n"
            "Zweiter Abschnitt.\n\n"
            "Erster Abschnitt mit Inhalt.\n\n"  # Duplikat
            "Dritter Abschnitt."
        )
        mock_response = MagicMock()
        mock_response.json.return_value = {"response": wiederholter_output, "eval_count": 50}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.llm import generate_text
            result = await generate_text("System", "User")

        assert result["text"].count("Erster Abschnitt") == 1
        assert "Zweiter Abschnitt." in result["text"]
        assert "Dritter Abschnitt." in result["text"]


# ══════════════════════════════════════════════════════════════════
# JOB QUEUE SERVICE
# ══════════════════════════════════════════════════════════════════

class TestJobQueueService:

    def setup_method(self):
        """Frische JobQueue-Instanz fuer jeden Test."""
        from app.services.job_queue import JobQueue
        self.queue = JobQueue()

    def test_create_job_felder(self):
        """Neu erstellter Job hat korrekte Felder und Status PENDING."""
        from app.services.job_queue import JobStatus
        job = self.queue.create_job("dokumentation", "Test")
        assert len(job.job_id) == 32
        assert job.workflow == "dokumentation"
        assert job.status == JobStatus.PENDING
        assert job.result_text is None
        assert job.error_msg is None
        assert job.started_at is None
        assert job.finished_at is None

    def test_to_dict_vollstaendig(self):
        """to_dict enthaelt alle erwarteten Schluessels."""
        job = self.queue.create_job("anamnese", "Test")
        d = job.to_dict()
        for key in ["job_id", "workflow", "status", "created_at",
                    "started_at", "finished_at", "result_text",
                    "error_msg", "duration_s", "model_used"]:
            assert key in d, f"Schluessel fehlt: {key}"

    @pytest.mark.asyncio
    async def test_run_job_done_pfad(self):
        """Erfolgreicher Job landet in Status DONE mit result_text."""
        from app.services.job_queue import JobStatus
        job = self.queue.create_job("dokumentation", "Test")

        async def _ok():
            return {"text": "Ergebnis", "model_used": "ollama/test"}

        await self.queue.run_job(job, _ok())

        assert job.status == JobStatus.DONE
        assert job.result_text == "Ergebnis"
        assert job.finished_at is not None
        assert job.duration_s is not None

    @pytest.mark.asyncio
    async def test_run_job_error_pfad(self):
        """Fehlerhafter Job landet in Status ERROR mit error_msg."""
        from app.services.job_queue import JobStatus
        job = self.queue.create_job("dokumentation", "Test")

        async def _fail():
            raise ValueError("Etwas ist schiefgelaufen")

        await self.queue.run_job(job, _fail())

        assert job.status == JobStatus.ERROR
        assert "schiefgelaufen" in job.error_msg
        assert job.finished_at is not None

    def test_cleanup_entfernt_alte_jobs(self):
        """Cleanup entfernt abgeschlossene Jobs wenn Limit ueberschritten."""
        from app.services.job_queue import JobQueue, JobStatus
        q = JobQueue()
        q._max_jobs = 5

        # 6 Jobs erstellen und alle auf DONE setzen
        for i in range(6):
            job = q.create_job("dokumentation", f"Job {i}")
            job.status = JobStatus.DONE

        # 7. Job triggert Cleanup
        q.create_job("dokumentation", "Trigger")

        assert len(q._jobs) <= 6

    def test_get_all_jobs_neueste_zuerst(self):
        """get_all_jobs gibt Jobs in absteigender Reihenfolge zurueck."""
        j1 = self.queue.create_job("dokumentation", "Erster")
        j2 = self.queue.create_job("anamnese", "Zweiter")
        j3 = self.queue.create_job("entlassbericht", "Dritter")

        jobs = self.queue.get_all_jobs()
        assert jobs[0].job_id == j3.job_id
        assert jobs[-1].job_id == j1.job_id

    def test_get_job_unbekannte_id(self):
        """Unbekannte Job-ID gibt None zurueck."""
        assert self.queue.get_job("nichtvorhanden") is None


# ══════════════════════════════════════════════════════════════════
# DOCX FILL SERVICE
# ══════════════════════════════════════════════════════════════════

class TestDocxFillService:

    def test_find_placeholders_doppelte_klammern(self):
        """{{FELDNAME}} wird korrekt erkannt."""
        from app.services.docx_fill import find_placeholders
        result = find_placeholders("Hallo {{NAME}}, dein Datum: {{DATUM}}")
        assert "NAME" in result
        assert "DATUM" in result

    def test_find_placeholders_eckige_klammern(self):
        """[FELDNAME] (nur Grossbuchstaben) wird korrekt erkannt."""
        from app.services.docx_fill import find_placeholders
        result = find_placeholders("Diagnose: [DIAGNOSE] Datum: [AUFNAHME DATUM]")
        assert "DIAGNOSE" in result
        assert "AUFNAHME DATUM" in result

    def test_find_placeholders_gemischt(self):
        """Beide Muster werden kombiniert erkannt."""
        from app.services.docx_fill import find_placeholders
        text = "{{PATIENT}} wurde am [AUFNAHMEDATUM] aufgenommen."
        result = find_placeholders(text)
        assert "PATIENT" in result
        assert "AUFNAHMEDATUM" in result

    def test_find_placeholders_keine_treffer(self):
        """Text ohne Platzhalter ergibt leere Liste."""
        from app.services.docx_fill import find_placeholders
        assert find_placeholders("Normaler Text ohne Platzhalter.") == []

    def test_find_placeholders_duplikate_werden_entfernt(self):
        """Gleicher Platzhalter mehrfach ergibt nur einen Eintrag."""
        from app.services.docx_fill import find_placeholders
        result = find_placeholders("{{NAME}} und nochmal {{NAME}}")
        assert result.count("NAME") == 1

    @pytest.mark.asyncio
    async def test_fill_docx_mit_platzhaltern(self, tmp_path):
        """Vorlage mit Platzhaltern wird korrekt befuellt."""
        from docx import Document
        from app.services.docx_fill import fill_docx_template

        # Vorlage mit Platzhalter erstellen
        doc = Document()
        doc.add_paragraph("Patient: {{NAME}}")
        template_path = tmp_path / "vorlage.docx"
        doc.save(str(template_path))

        generated = "NAME: Max Mustermann\nDiagnose: F32.1"
        result_path = await fill_docx_template(
            template_path, "Verlauf...", generated, tmp_path, "test"
        )

        assert result_path.exists()
        result_doc = Document(str(result_path))
        full_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "{{NAME}}" not in full_text

    @pytest.mark.asyncio
    async def test_fill_docx_ohne_platzhalter_append(self, tmp_path):
        """Vorlage ohne Platzhalter bekommt generierten Text angehaengt."""
        from docx import Document
        from app.services.docx_fill import fill_docx_template

        doc = Document()
        doc.add_paragraph("Bestehender Inhalt ohne Platzhalter.")
        template_path = tmp_path / "vorlage.docx"
        doc.save(str(template_path))

        result_path = await fill_docx_template(
            template_path, "Verlauf", "# Befund\nPatient stabil.", tmp_path, "test"
        )

        assert result_path.exists()
        result_doc = Document(str(result_path))
        full_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "Generierter Inhalt" in full_text


# ══════════════════════════════════════════════════════════════════
# EMBEDDINGS SERVICE
# ══════════════════════════════════════════════════════════════════

class TestEmbeddingsService:

    @pytest.mark.asyncio
    async def test_get_embedding_erfolgreich(self):
        """Erfolgreiches Embedding gibt Liste mit 768 Floats zurueck."""
        fake_vec = [0.1] * 768
        mock_response = MagicMock()
        mock_response.json.return_value = {"embedding": fake_vec}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is not None
        assert len(result) == 768

    @pytest.mark.asyncio
    async def test_get_embedding_connect_error_gibt_none(self):
        """ConnectError gibt None zurueck (kein Absturz)."""
        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                side_effect=httpx.ConnectError("refused")
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is None

    @pytest.mark.asyncio
    async def test_get_embedding_falsche_dimension_gibt_none(self):
        """Embedding mit falscher Dimension gibt None zurueck."""
        mock_response = MagicMock()
        mock_response.json.return_value = {"embedding": [0.1] * 512}  # falsch: 512 statt 768
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is None

    @pytest.mark.asyncio
    async def test_get_embedding_leere_response_gibt_none(self):
        """Fehlende 'embedding'-Key in Response gibt None zurueck."""
        mock_response = MagicMock()
        mock_response.json.return_value = {}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is None

    @pytest.mark.asyncio
    async def test_retrieve_style_examples_leer_wenn_keine_eintraege(self, init_test_db):
        """Kein Eintrag in DB gibt leeren String zurueck."""
        from sqlalchemy.ext.asyncio import AsyncSession
        from app.core.database import engine
        from app.services.embeddings import retrieve_style_examples

        # pgvector-Operator <=> nicht in SQLite verfuegbar –
        # semantische Suche wird gemockt, nur der Fallback-Pfad getestet
        with patch("app.services.embeddings.get_embedding", new=AsyncMock(return_value=None)):
            async with AsyncSession(engine) as db:
                result = await retrieve_style_examples(
                    db, "unbekannter_therapeut", "dokumentation", "Testtext"
                )

        assert result == ""

    @pytest.mark.asyncio
    async def test_retrieve_style_examples_fallback_ohne_embedding(self, init_test_db):
        """Wenn get_embedding None zurueckgibt, wird Fallback (neueste) verwendet."""
        from sqlalchemy.ext.asyncio import AsyncSession
        from app.core.database import engine
        from app.models.db import StyleEmbedding
        from app.services.embeddings import retrieve_style_examples
        from datetime import datetime, timezone

        # Beispiel in DB schreiben
        async with AsyncSession(engine) as db:
            example = StyleEmbedding(
                id=uuid.uuid4().hex,
                therapeut_id="test_therapeut",
                dokumenttyp="dokumentation",
                text="Beispieltext für Fallback-Test.",
                word_count=5,
                ist_statisch=False,
                created_at=datetime.now(timezone.utc),
            )
            db.add(example)
            await db.commit()

        with patch("app.services.embeddings.get_embedding", new=AsyncMock(return_value=None)):
            async with AsyncSession(engine) as db:
                result = await retrieve_style_examples(
                    db, "test_therapeut", "dokumentation", "Testtext"
                )

        assert "Beispieltext" in result

    @pytest.mark.asyncio
    async def test_retrieve_style_examples_statische_anker_immer_dabei(self, init_test_db):
        """Statische Anker-Beispiele werden immer eingeschlossen."""
        from sqlalchemy.ext.asyncio import AsyncSession
        from app.core.database import engine
        from app.models.db import StyleEmbedding
        from app.services.embeddings import retrieve_style_examples
        from datetime import datetime, timezone

        async with AsyncSession(engine) as db:
            anker = StyleEmbedding(
                id=uuid.uuid4().hex,
                therapeut_id="anker_therapeut",
                dokumenttyp="dokumentation",
                text="Dies ist ein Anker-Beispiel.",
                word_count=5,
                ist_statisch=True,
                created_at=datetime.now(timezone.utc),
            )
            db.add(anker)
            await db.commit()

        with patch("app.services.embeddings.get_embedding", new=AsyncMock(return_value=None)):
            async with AsyncSession(engine) as db:
                result = await retrieve_style_examples(
                    db, "anker_therapeut", "dokumentation", ""
                )

        assert "Anker" in result
        assert "Anker-Beispiel" in result


# ══════════════════════════════════════════════════════════════════
# SPRECHER-ANNOTATION (Unit Tests)
# Testet beide Annotationswege:
#   1. pyannote-basiert: _assign_speaker_from_diarization()
#   2. Pausen-Heuristik: _assign_speakers()
# Kein Whisper, kein pyannote, keine GPU nötig.
# ══════════════════════════════════════════════════════════════════

class TestSpreecherAnnotation:
    """Tests für Sprecher-Zuweisung – beide Implementierungen."""

    # ── Hilfsmethoden ─────────────────────────────────────────────

    def _seg(self, start, end, text):
        """Erstellt ein Minimal-Segment-Objekt (Whisper-kompatibel)."""
        class Seg:
            pass
        s = Seg()
        s.start = start
        s.end = end
        s.text = text
        return s

    def _dia(self, start, end, speaker):
        """Erstellt ein pyannote-Diarization-Segment."""
        return {"start": start, "end": end, "speaker": speaker}

    # ── _assign_speaker_from_diarization ──────────────────────────

    def test_exakte_ueberlappung(self):
        """Segment liegt vollständig in einem Diarization-Segment."""
        from app.services.transcription import _assign_speaker_from_diarization
        dia = [self._dia(0.0, 5.0, "A"), self._dia(5.0, 10.0, "B")]
        assert _assign_speaker_from_diarization(1.0, 3.0, dia) == "A"
        assert _assign_speaker_from_diarization(6.0, 8.0, dia) == "B"

    def test_groesste_ueberlappung_gewinnt(self):
        """Bei Überschneidung gewinnt der Sprecher mit mehr Überlappung."""
        from app.services.transcription import _assign_speaker_from_diarization
        # Segment 4.0-6.0 überlappt mit A (4.0-5.0 = 1s) und B (5.0-6.0 = 1s)
        # Gleiche Überlappung → A gewinnt (erster Fund)
        dia = [self._dia(0.0, 5.0, "A"), self._dia(5.0, 10.0, "B")]
        result = _assign_speaker_from_diarization(4.0, 6.0, dia)
        assert result in ("A", "B")  # deterministisch aber beide valide

    def test_klare_mehrheit(self):
        """Segment überlappt stark mit einem Sprecher."""
        from app.services.transcription import _assign_speaker_from_diarization
        dia = [self._dia(0.0, 5.0, "A"), self._dia(5.0, 10.0, "B")]
        # 4.5-7.0: 0.5s mit A, 2.0s mit B → B gewinnt
        assert _assign_speaker_from_diarization(4.5, 7.0, dia) == "B"

    def test_keine_ueberlappung_gibt_none(self):
        """Segment liegt außerhalb aller Diarization-Segmente → None."""
        from app.services.transcription import _assign_speaker_from_diarization
        dia = [self._dia(0.0, 2.0, "A"), self._dia(3.0, 5.0, "B")]
        # Lücke zwischen 2.0 und 3.0
        assert _assign_speaker_from_diarization(2.1, 2.9, dia) is None

    def test_leere_diarization_gibt_none(self):
        """Leere Diarization-Liste gibt None zurück."""
        from app.services.transcription import _assign_speaker_from_diarization
        assert _assign_speaker_from_diarization(0.0, 1.0, []) is None

    def test_drei_sprecher(self):
        """Auch mehr als zwei Sprecher werden korrekt zugewiesen."""
        from app.services.transcription import _assign_speaker_from_diarization
        dia = [
            self._dia(0.0, 3.0, "A"),
            self._dia(3.0, 6.0, "B"),
            self._dia(6.0, 9.0, "C"),
        ]
        assert _assign_speaker_from_diarization(0.5, 1.5, dia) == "A"
        assert _assign_speaker_from_diarization(3.5, 4.5, dia) == "B"
        assert _assign_speaker_from_diarization(6.5, 7.5, dia) == "C"

    def test_punkt_ueberlappung_zaehlt_nicht(self):
        """Berührung ohne echte Überlappung (overlap=0) gibt None."""
        from app.services.transcription import _assign_speaker_from_diarization
        dia = [self._dia(0.0, 5.0, "A")]
        # Segment beginnt genau wo Diarization endet → overlap=0
        assert _assign_speaker_from_diarization(5.0, 6.0, dia) is None

    def test_realistische_therapiesitzung(self):
        """Simuliert typisches Therapiegespräch mit abwechselnden Sprechern."""
        from app.services.transcription import _assign_speaker_from_diarization
        # Therapeut: 0-8s, 15-22s, 30-35s
        # Klient:    8-15s, 22-30s, 35-50s
        dia = [
            self._dia(0.0,  8.0,  "A"),   # Therapeut
            self._dia(8.0,  15.0, "B"),   # Klient
            self._dia(15.0, 22.0, "A"),   # Therapeut
            self._dia(22.0, 30.0, "B"),   # Klient
            self._dia(30.0, 35.0, "A"),   # Therapeut
            self._dia(35.0, 50.0, "B"),   # Klient (langer Monolog)
        ]
        whisper_segs = [
            (1.0,  4.0,  "A"),   # Therapeut-Frage
            (9.0,  13.0, "B"),   # Klient-Antwort
            (16.0, 20.0, "A"),   # Therapeut-Intervention
            (23.0, 28.0, "B"),   # Klient erzählt
            (31.0, 34.0, "A"),   # Therapeut kurz
            (36.0, 45.0, "B"),   # Klient langer Bericht
        ]
        for start, end, expected in whisper_segs:
            result = _assign_speaker_from_diarization(start, end, dia)
            assert result == expected, \
                f"Segment {start}-{end}s: erwartet {expected}, bekommen {result}"

    # ── _assign_speakers (Pausen-Heuristik) ───────────────────────

    def test_heuristik_wechsel_bei_langer_pause(self):
        """Pause ≥ 1.2s → Sprecherwechsel."""
        from app.services.transcription import _assign_speakers
        segs = [
            self._seg(0.0, 3.0, "Guten Morgen."),
            self._seg(4.5, 7.0, "Wie geht es Ihnen?"),  # 1.5s Pause → Wechsel
        ]
        result = _assign_speakers(segs)
        assert "[A]: Guten Morgen." in result
        assert "[B]: Wie geht es Ihnen?" in result

    def test_heuristik_kein_wechsel_bei_kurzer_pause(self):
        """Pause < 1.2s → kein Sprecherwechsel."""
        from app.services.transcription import _assign_speakers
        segs = [
            self._seg(0.0, 2.0, "Erster Satz."),
            self._seg(2.5, 4.0, "Zweiter Satz."),  # 0.5s Pause → kein Wechsel
        ]
        result = _assign_speakers(segs)
        lines = result.strip().split("\n")
        assert all(l.startswith("[A]:") for l in lines), \
            f"Beide Segmente sollten [A] sein: {result}"

    def test_heuristik_leere_segmente_werden_uebersprungen(self):
        """Segmente mit leerem Text werden nicht in den Output aufgenommen."""
        from app.services.transcription import _assign_speakers
        segs = [
            self._seg(0.0, 1.0, ""),
            self._seg(1.0, 2.0, "  "),
            self._seg(2.0, 3.0, "Echter Text."),
        ]
        result = _assign_speakers(segs)
        assert result.count("\n") == 0  # nur eine Zeile
        assert "Echter Text." in result

    def test_heuristik_alternierende_sprecher(self):
        """Typisches Therapiegespräch mit regelmäßigem Wechsel."""
        from app.services.transcription import _assign_speakers
        segs = [
            self._seg(0.0,  3.0, "Frage des Therapeuten."),
            self._seg(4.5,  8.0, "Antwort des Klienten."),    # Wechsel → B
            self._seg(9.5, 11.0, "Nachfrage Therapeut."),     # Wechsel → A
            self._seg(12.5, 16.0, "Weitere Antwort Klient."), # Wechsel → B
        ]
        result = _assign_speakers(segs)
        lines = result.strip().split("\n")
        assert lines[0].startswith("[A]:")
        assert lines[1].startswith("[B]:")
        assert lines[2].startswith("[A]:")
        assert lines[3].startswith("[B]:")

    def test_heuristik_erster_sprecher_immer_a(self):
        """Das erste Segment wird immer [A] zugewiesen."""
        from app.services.transcription import _assign_speakers
        segs = [self._seg(5.0, 8.0, "Erster Sprecher.")]
        assert _assign_speakers(segs).startswith("[A]:")

    def test_heuristik_genau_an_schwelle(self):
        """Pause exakt 1.2s → Wechsel (Schwellenwert inklusiv)."""
        from app.services.transcription import _assign_speakers
        segs = [
            self._seg(0.0, 2.0, "Erster."),
            self._seg(3.2, 5.0, "Zweiter."),  # genau 1.2s Pause
        ]
        result = _assign_speakers(segs)
        assert "[B]:" in result

    # ── _diarize Fallback-Verhalten ────────────────────────────────

    def test_diarize_gibt_none_wenn_deaktiviert(self):
        """_diarize() gibt None zurück wenn DIARIZATION_ENABLED=false."""
        from app.services import transcription as t
        from unittest.mock import patch
        with patch.object(t.settings, "DIARIZATION_ENABLED", False):
            result = t._diarize(Path("/tmp/test.wav"))
        assert result is None

    def test_diarize_gibt_none_wenn_pyannote_fehlt(self):
        """_diarize() gibt None zurück wenn pyannote nicht installiert."""
        from app.services import transcription as t
        from unittest.mock import patch
        import sys
        with patch.object(t.settings, "DIARIZATION_ENABLED", True), \
             patch.object(t, "_diarization_pipeline", None), \
             patch.dict(sys.modules, {"pyannote.audio": None}):
            result = t._diarize(Path("/tmp/test.wav"))
        assert result is None

    def test_get_diarization_pipeline_gibt_none_wenn_deaktiviert(self):
        """_get_diarization_pipeline() gibt None wenn DIARIZATION_ENABLED=false."""
        from app.services import transcription as t
        from unittest.mock import patch
        with patch.object(t.settings, "DIARIZATION_ENABLED", False):
            result = t._get_diarization_pipeline()
        assert result is None
