"""
Tests fuer Service-Schicht: llm, job_queue, docx_fill, embeddings.

Fokus auf bisher unabgedeckte Pfade:
  - llm.py:       Timeout, HTTPStatusError, leere Response
  - job_queue.py: Error-Pfad, Cleanup, to_dict-Felder
  - docx_fill.py: find_placeholders, Platzhalter-Ersetzung, Append-Pfad
  - embeddings.py: get_embedding Fehlerfaelle, retrieve_style_examples Fallback
"""
import asyncio
import uuid
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import httpx
import pytest

# ══════════════════════════════════════════════════════════════════
# LLM SERVICE
# ══════════════════════════════════════════════════════════════════

class TestLLMService:

    @pytest.mark.asyncio
    async def test_generate_text_erfolgreich(self):
        """generate_text gibt text, model_used und duration_s zurueck."""
        mock_response = MagicMock()
        mock_response.json.return_value = {
            "response": "Generierter Text.",
            "eval_count": 42,
        }
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.llm import generate_text
            result = await generate_text("System", "User")

        assert result["text"] == "Generierter Text."
        assert "ollama/" in result["model_used"]
        assert result["token_count"] == 42
        assert "duration_s" in result

    @pytest.mark.asyncio
    async def test_generate_text_connect_error_wirft_runtime(self):
        """ConnectError wird als RuntimeError mit sprechender Meldung weitergegeben."""
        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                side_effect=httpx.ConnectError("refused")
            )
            from app.services.llm import generate_text
            with pytest.raises(RuntimeError, match="Ollama nicht erreichbar"):
                await generate_text("System", "User")

    @pytest.mark.asyncio
    async def test_generate_text_http_status_error(self):
        """HTTPStatusError (z.B. 500) wird als RuntimeError weitergegeben."""
        mock_response = MagicMock()
        mock_response.status_code = 500
        mock_response.text = "Internal Server Error"

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                side_effect=httpx.HTTPStatusError(
                    "500", request=MagicMock(), response=mock_response
                )
            )
            from app.services.llm import generate_text
            with pytest.raises(RuntimeError, match="Ollama Fehler 500"):
                await generate_text("System", "User")

    @pytest.mark.asyncio
    async def test_generate_text_leere_response(self):
        """Leere Ollama-Response liefert leeren Text ohne Fehler."""
        mock_response = MagicMock()
        mock_response.json.return_value = {"response": ""}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.llm import generate_text
            result = await generate_text("System", "User")

        assert result["text"] == ""


# ══════════════════════════════════════════════════════════════════
# JOB QUEUE SERVICE
# ══════════════════════════════════════════════════════════════════

class TestJobQueueService:

    def setup_method(self):
        """Frische JobQueue-Instanz fuer jeden Test."""
        from app.services.job_queue import JobQueue
        self.queue = JobQueue()

    def test_create_job_felder(self):
        """Neu erstellter Job hat korrekte Felder und Status PENDING."""
        from app.services.job_queue import JobStatus
        job = self.queue.create_job("dokumentation", "Test")
        assert len(job.job_id) == 32
        assert job.workflow == "dokumentation"
        assert job.status == JobStatus.PENDING
        assert job.result_text is None
        assert job.error_msg is None
        assert job.started_at is None
        assert job.finished_at is None

    def test_to_dict_vollstaendig(self):
        """to_dict enthaelt alle erwarteten Schluessels."""
        job = self.queue.create_job("anamnese", "Test")
        d = job.to_dict()
        for key in ["job_id", "workflow", "status", "created_at",
                    "started_at", "finished_at", "result_text",
                    "error_msg", "duration_s", "model_used"]:
            assert key in d, f"Schluessel fehlt: {key}"

    @pytest.mark.asyncio
    async def test_run_job_done_pfad(self):
        """Erfolgreicher Job landet in Status DONE mit result_text."""
        from app.services.job_queue import JobStatus
        job = self.queue.create_job("dokumentation", "Test")

        async def _ok():
            return {"text": "Ergebnis", "model_used": "ollama/test"}

        await self.queue.run_job(job, _ok())

        assert job.status == JobStatus.DONE
        assert job.result_text == "Ergebnis"
        assert job.finished_at is not None
        assert job.duration_s is not None

    @pytest.mark.asyncio
    async def test_run_job_error_pfad(self):
        """Fehlerhafter Job landet in Status ERROR mit error_msg."""
        from app.services.job_queue import JobStatus
        job = self.queue.create_job("dokumentation", "Test")

        async def _fail():
            raise ValueError("Etwas ist schiefgelaufen")

        await self.queue.run_job(job, _fail())

        assert job.status == JobStatus.ERROR
        assert "schiefgelaufen" in job.error_msg
        assert job.finished_at is not None

    def test_cleanup_entfernt_alte_jobs(self):
        """Cleanup entfernt abgeschlossene Jobs wenn Limit ueberschritten."""
        from app.services.job_queue import JobQueue, JobStatus
        q = JobQueue()
        q._max_jobs = 5

        # 6 Jobs erstellen und alle auf DONE setzen
        for i in range(6):
            job = q.create_job("dokumentation", f"Job {i}")
            job.status = JobStatus.DONE

        # 7. Job triggert Cleanup
        q.create_job("dokumentation", "Trigger")

        assert len(q._jobs) <= 6

    def test_get_all_jobs_neueste_zuerst(self):
        """get_all_jobs gibt Jobs in absteigender Reihenfolge zurueck."""
        j1 = self.queue.create_job("dokumentation", "Erster")
        j2 = self.queue.create_job("anamnese", "Zweiter")
        j3 = self.queue.create_job("entlassbericht", "Dritter")

        jobs = self.queue.get_all_jobs()
        assert jobs[0].job_id == j3.job_id
        assert jobs[-1].job_id == j1.job_id

    def test_get_job_unbekannte_id(self):
        """Unbekannte Job-ID gibt None zurueck."""
        assert self.queue.get_job("nichtvorhanden") is None


# ══════════════════════════════════════════════════════════════════
# DOCX FILL SERVICE
# ══════════════════════════════════════════════════════════════════

class TestDocxFillService:

    def test_find_placeholders_doppelte_klammern(self):
        """{{FELDNAME}} wird korrekt erkannt."""
        from app.services.docx_fill import find_placeholders
        result = find_placeholders("Hallo {{NAME}}, dein Datum: {{DATUM}}")
        assert "NAME" in result
        assert "DATUM" in result

    def test_find_placeholders_eckige_klammern(self):
        """[FELDNAME] (nur Grossbuchstaben) wird korrekt erkannt."""
        from app.services.docx_fill import find_placeholders
        result = find_placeholders("Diagnose: [DIAGNOSE] Datum: [AUFNAHME DATUM]")
        assert "DIAGNOSE" in result
        assert "AUFNAHME DATUM" in result

    def test_find_placeholders_gemischt(self):
        """Beide Muster werden kombiniert erkannt."""
        from app.services.docx_fill import find_placeholders
        text = "{{PATIENT}} wurde am [AUFNAHMEDATUM] aufgenommen."
        result = find_placeholders(text)
        assert "PATIENT" in result
        assert "AUFNAHMEDATUM" in result

    def test_find_placeholders_keine_treffer(self):
        """Text ohne Platzhalter ergibt leere Liste."""
        from app.services.docx_fill import find_placeholders
        assert find_placeholders("Normaler Text ohne Platzhalter.") == []

    def test_find_placeholders_duplikate_werden_entfernt(self):
        """Gleicher Platzhalter mehrfach ergibt nur einen Eintrag."""
        from app.services.docx_fill import find_placeholders
        result = find_placeholders("{{NAME}} und nochmal {{NAME}}")
        assert result.count("NAME") == 1

    @pytest.mark.asyncio
    async def test_fill_docx_mit_platzhaltern(self, tmp_path):
        """Vorlage mit Platzhaltern wird korrekt befuellt."""
        from docx import Document
        from app.services.docx_fill import fill_docx_template

        # Vorlage mit Platzhalter erstellen
        doc = Document()
        doc.add_paragraph("Patient: {{NAME}}")
        template_path = tmp_path / "vorlage.docx"
        doc.save(str(template_path))

        generated = "NAME: Max Mustermann\nDiagnose: F32.1"
        result_path = await fill_docx_template(
            template_path, "Verlauf...", generated, tmp_path, "test"
        )

        assert result_path.exists()
        result_doc = Document(str(result_path))
        full_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "{{NAME}}" not in full_text

    @pytest.mark.asyncio
    async def test_fill_docx_ohne_platzhalter_append(self, tmp_path):
        """Vorlage ohne Platzhalter bekommt generierten Text angehaengt."""
        from docx import Document
        from app.services.docx_fill import fill_docx_template

        doc = Document()
        doc.add_paragraph("Bestehender Inhalt ohne Platzhalter.")
        template_path = tmp_path / "vorlage.docx"
        doc.save(str(template_path))

        result_path = await fill_docx_template(
            template_path, "Verlauf", "# Befund\nPatient stabil.", tmp_path, "test"
        )

        assert result_path.exists()
        result_doc = Document(str(result_path))
        full_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "Generierter Inhalt" in full_text


# ══════════════════════════════════════════════════════════════════
# EMBEDDINGS SERVICE
# ══════════════════════════════════════════════════════════════════

class TestEmbeddingsService:

    @pytest.mark.asyncio
    async def test_get_embedding_erfolgreich(self):
        """Erfolgreiches Embedding gibt Liste mit 768 Floats zurueck."""
        fake_vec = [0.1] * 768
        mock_response = MagicMock()
        mock_response.json.return_value = {"embedding": fake_vec}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is not None
        assert len(result) == 768

    @pytest.mark.asyncio
    async def test_get_embedding_connect_error_gibt_none(self):
        """ConnectError gibt None zurueck (kein Absturz)."""
        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                side_effect=httpx.ConnectError("refused")
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is None

    @pytest.mark.asyncio
    async def test_get_embedding_falsche_dimension_gibt_none(self):
        """Embedding mit falscher Dimension gibt None zurueck."""
        mock_response = MagicMock()
        mock_response.json.return_value = {"embedding": [0.1] * 512}  # falsch: 512 statt 768
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is None

    @pytest.mark.asyncio
    async def test_get_embedding_leere_response_gibt_none(self):
        """Fehlende 'embedding'-Key in Response gibt None zurueck."""
        mock_response = MagicMock()
        mock_response.json.return_value = {}
        mock_response.raise_for_status = MagicMock()

        with patch("httpx.AsyncClient") as mock_client:
            mock_client.return_value.__aenter__.return_value.post = AsyncMock(
                return_value=mock_response
            )
            from app.services.embeddings import get_embedding
            result = await get_embedding("Testtext")

        assert result is None

    @pytest.mark.asyncio
    async def test_retrieve_style_examples_leer_wenn_keine_eintraege(self, init_test_db):
        """Kein Eintrag in DB gibt leeren String zurueck."""
        from sqlalchemy.ext.asyncio import AsyncSession
        from app.core.database import engine
        from app.services.embeddings import retrieve_style_examples

        async with AsyncSession(engine) as db:
            result = await retrieve_style_examples(
                db, "unbekannter_therapeut", "dokumentation", "Testtext"
            )

        assert result == ""

    @pytest.mark.asyncio
    async def test_retrieve_style_examples_fallback_ohne_embedding(self, init_test_db):
        """Wenn get_embedding None zurueckgibt, wird Fallback (neueste) verwendet."""
        from sqlalchemy.ext.asyncio import AsyncSession
        from app.core.database import engine
        from app.models.db import StyleEmbedding
        from app.services.embeddings import retrieve_style_examples
        from datetime import datetime, timezone

        # Beispiel in DB schreiben
        async with AsyncSession(engine) as db:
            example = StyleEmbedding(
                id=uuid.uuid4().hex,
                therapeut_id="test_therapeut",
                dokumenttyp="dokumentation",
                text="Beispieltext für Fallback-Test.",
                word_count=5,
                ist_statisch=False,
                created_at=datetime.now(timezone.utc),
            )
            db.add(example)
            await db.commit()

        with patch("app.services.embeddings.get_embedding", new=AsyncMock(return_value=None)):
            async with AsyncSession(engine) as db:
                result = await retrieve_style_examples(
                    db, "test_therapeut", "dokumentation", "Testtext"
                )

        assert "Beispieltext" in result

    @pytest.mark.asyncio
    async def test_retrieve_style_examples_statische_anker_immer_dabei(self, init_test_db):
        """Statische Anker-Beispiele werden immer eingeschlossen."""
        from sqlalchemy.ext.asyncio import AsyncSession
        from app.core.database import engine
        from app.models.db import StyleEmbedding
        from app.services.embeddings import retrieve_style_examples
        from datetime import datetime, timezone

        async with AsyncSession(engine) as db:
            anker = StyleEmbedding(
                id=uuid.uuid4().hex,
                therapeut_id="anker_therapeut",
                dokumenttyp="dokumentation",
                text="Dies ist ein Anker-Beispiel.",
                word_count=5,
                ist_statisch=True,
                created_at=datetime.now(timezone.utc),
            )
            db.add(anker)
            await db.commit()

        with patch("app.services.embeddings.get_embedding", new=AsyncMock(return_value=None)):
            async with AsyncSession(engine) as db:
                result = await retrieve_style_examples(
                    db, "anker_therapeut", "dokumentation", ""
                )

        assert "Anker" in result
        assert "Anker-Beispiel" in result
