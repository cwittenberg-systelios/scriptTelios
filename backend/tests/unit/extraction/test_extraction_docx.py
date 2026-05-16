"""
Unit-Tests fuer extract_docx_section in extraction.py.

Erstellt synthetische DOCX-Dateien mit verschiedenen Strukturen
(Heading-Style, Bold-Heading, Plain-Text-Heading) und prueft
ob der richtige Abschnitt extrahiert wird.

Hintergrund: Bug aus Session 2026-04-22 — bidirektionaler Substring-Match
war zu großzügig: 'AKUTAUFNAHME' matchte 'Begründung für die Akutaufnahme'
als Teilstring. Die Tests stellen sicher dass die Heading-Erkennung praezise
ist.
"""
import pytest
from pathlib import Path

# python-docx fuer DOCX-Erstellung in den Tests
docx = pytest.importorskip("docx", reason="python-docx fehlt")
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.services.extraction import extract_docx_section, STYLE_SECTION_HEADINGS


# ─────────────────────────────────────────────────────────────────────────────
# Helper: DOCX mit Headings/Bold-Headings konstruieren
# ─────────────────────────────────────────────────────────────────────────────


def _make_docx_with_headings(tmp_path: Path, sections: list[tuple[str, str, str]]) -> Path:
    """
    Erstellt ein DOCX mit den gegebenen Abschnitten.

    sections: Liste von (heading_text, body_text, style)
              style ∈ {"heading", "bold", "plain"}
    """
    doc = Document()
    for heading, body, style in sections:
        if style == "heading":
            doc.add_heading(heading, level=2)
        elif style == "bold":
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.bold = True
        else:  # plain
            doc.add_paragraph(heading)
        # Body
        if body:
            doc.add_paragraph(body)
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Heading-Style Erkennung
# ─────────────────────────────────────────────────────────────────────────────


class TestExtractDocxSectionHeading:
    """Heading-Erkennung via DOCX-Heading-Style."""

    def test_entlassbericht_findet_psychotherapeutischen_verlauf(self, tmp_path):
        body = (
            "Zu Beginn des stationaeren Aufenthaltes formulierte die Patientin "
            "als zentrales Anliegen wieder inneren Halt zu finden. Wir erlebten "
            "sie zu Therapiebeginn deutlich erschoepft und in ihrem Selbstwert "
            "erheblich verunsichert. Im Verlauf zeigte sich eine deutliche "
            "Stabilisierung. Die Therapieziele konnten teilweise erreicht werden. "
            "Hypnosystemische Anteilearbeit wurde erfolgreich eingesetzt. "
            "Die Patientin wurde am Ende stabil entlassen."
        )
        path = _make_docx_with_headings(tmp_path, [
            ("Anamnese", "Kurzer Anamnese-Abschnitt der nicht relevant ist.", "heading"),
            ("Psychotherapeutischer Verlauf", body, "heading"),
            ("Diagnosen", "F33.2", "heading"),
        ])
        result = extract_docx_section(path, "entlassbericht")
        assert "Hypnosystemische Anteilearbeit" in result
        assert "F33.2" not in result  # naechster Heading beendet Section
        assert "Anamnese-Abschnitt" not in result  # vorheriger Abschnitt nicht enthalten

    def test_anamnese_findet_aktuelle_anamnese(self, tmp_path):
        body = (
            "Die Patientin stellt sich mit dem Hauptanliegen vor sich in ihrer "
            "aktuellen psychischen Situation besser verstehen zu koennen. Sie "
            "beschreibt zunehmende Erschoepfung und innere Unruhe. Die Symptome "
            "haben sich schleichend entwickelt. Eine traumatische Vorgeschichte "
            "ist aus den Quellen bekannt."
        )
        path = _make_docx_with_headings(tmp_path, [
            ("Stammdaten", "Geburt: 1.1.1985", "heading"),
            ("Aktuelle Anamnese", body, "heading"),
        ])
        result = extract_docx_section(path, "anamnese")
        assert "Hauptanliegen" in result
        assert "Geburt" not in result


# ─────────────────────────────────────────────────────────────────────────────
# Bold-Text als Heading erkennen (kein Heading-Style)
# ─────────────────────────────────────────────────────────────────────────────


class TestExtractDocxSectionBold:
    """Bold-Text Erkennung als Heading."""

    def test_bold_heading_wird_als_heading_erkannt(self, tmp_path):
        body = (
            "Im bisherigen Verlauf der Behandlung zeigte sich eine deutliche "
            "Verbesserung. Die Patientin hat ihre Selbstregulation deutlich "
            "ausgebaut und kann nun besser mit Stress umgehen. Allerdings "
            "ist eine ambulante Stabilisierung noch nicht moeglich. "
            "Wir empfehlen daher eine weitere stationaere Behandlung um die "
            "begonnene Traumabearbeitung fortzusetzen."
        )
        path = _make_docx_with_headings(tmp_path, [
            ("Bisheriger Verlauf", body, "bold"),  # NICHT Heading-Style sondern Bold
        ])
        result = extract_docx_section(path, "verlaengerung")
        assert "Selbstregulation" in result
        assert "Traumabearbeitung" in result


# ─────────────────────────────────────────────────────────────────────────────
# Bug-Regression: Substring-Match darf nicht zu groszuegig sein
# ─────────────────────────────────────────────────────────────────────────────


class TestExtractDocxSectionSubstringRegression:
    """
    Regression-Test fuer Bug aus 2026-04-22:
    'AKUTAUFNAHME' (Briefkopf-Heading) matchte faelschlich
    'Begründung für die Akutaufnahme' als Teilstring.
    """

    def test_kurzes_briefkopf_heading_matcht_nicht_langes_workflow_heading(self, tmp_path):
        # Briefkopf "AKUTAUFNAHME" ist kurz (12 Zeichen)
        # Das echte Heading "Begründung für Akutaufnahme" ist lang
        # Bei kurzen Headings (< 20 Zeichen) wird nur exakt/startswith gematcht
        # → "AKUTAUFNAHME" allein darf NICHT als "Begründung" Heading interpretiert werden
        body_kurz = "Briefkopf-Inhalt"
        body_lang = (
            "Die akute stationaere Behandlung ist medizinisch unabdingbar weil die "
            "Patientin in einer akuten Suizidkrise ist. Eine ambulante Versorgung "
            "ist aufgrund der Schwere nicht ausreichend. Die Symptomatik hat sich "
            "in den letzten Wochen deutlich verschlechtert. Eine sofortige stationaere "
            "Aufnahme ist erforderlich um eine weitere Verschlechterung zu verhindern."
        )
        path = _make_docx_with_headings(tmp_path, [
            ("AKUTAUFNAHME", body_kurz, "bold"),  # Briefkopf
            ("Begründung für Akutaufnahme", body_lang, "bold"),  # echte Section
        ])
        result = extract_docx_section(path, "akutantrag")
        # Erwartet: Inhalt der echten Section, nicht "Briefkopf-Inhalt"
        assert "akute stationaere Behandlung" in result
        assert "Briefkopf-Inhalt" not in result


# ─────────────────────────────────────────────────────────────────────────────
# Akutantrag-Fallback: Aktuelle Anamnese wenn Begruendung leer
# ─────────────────────────────────────────────────────────────────────────────


class TestExtractDocxSectionAkutFallback:
    """
    Test fuer den Fallback: wenn alle 'Begruendung' Sections leer sind,
    wird 'Aktuelle Anamnese' als Fallback verwendet (laut STYLE_SECTION_HEADINGS).
    """

    def test_akutantrag_fallback_auf_aktuelle_anamnese(self, tmp_path):
        body_anamnese = (
            "Die Patientin berichtet von einer schweren depressiven Episode mit "
            "Schlafstoerungen und Antriebslosigkeit. Sie kommt morgens kaum aus "
            "dem Bett und kann ihren Beruf nicht mehr ausueben. In den letzten "
            "Wochen sind erstmals lebensmuede Gedanken aufgetreten. Eine ambulante "
            "Therapie wurde bereits versucht aber als nicht ausreichend erlebt."
        )
        path = _make_docx_with_headings(tmp_path, [
            ("Aktuelle Anamnese", body_anamnese, "bold"),
            # KEINE Begruendung-Section vorhanden
        ])
        result = extract_docx_section(path, "akutantrag")
        assert "schweren depressiven Episode" in result


# ─────────────────────────────────────────────────────────────────────────────
# Edge Cases
# ─────────────────────────────────────────────────────────────────────────────


class TestExtractDocxSectionEdgeCases:

    def test_unbekannter_workflow_gibt_volltext_zurueck(self, tmp_path):
        path = _make_docx_with_headings(tmp_path, [
            ("Heading 1", "Body 1", "heading"),
            ("Heading 2", "Body 2", "heading"),
        ])
        # Workflow gibt es nicht in STYLE_SECTION_HEADINGS
        result = extract_docx_section(path, "unknown_workflow")
        # Fallback: Volltext
        assert "Heading 1" in result or "Body 1" in result

    def test_leeres_docx_gibt_leerstring(self, tmp_path):
        doc = Document()
        path = tmp_path / "leer.docx"
        doc.save(str(path))
        result = extract_docx_section(path, "entlassbericht")
        # Sollte nicht crashen
        assert isinstance(result, str)

    def test_alle_workflows_haben_headings_definiert(self):
        # Sanity-Check: alle bekannten Workflows haben Headings
        expected_workflows = {
            "entlassbericht", "verlaengerung", "folgeverlaengerung",
            "anamnese", "dokumentation", "akutantrag",
        }
        for workflow in expected_workflows:
            assert workflow in STYLE_SECTION_HEADINGS, f"Fehlt: {workflow}"
            assert len(STYLE_SECTION_HEADINGS[workflow]) > 0
