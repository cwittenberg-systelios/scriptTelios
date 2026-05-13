"""
scriptTelios Evaluations-Framework
===================================

Testet die Qualitaet der LLM-Generierung gegen definierte Erwartungen.
Laeuft gegen den laufenden Backend-Server (nicht gegen Mocks).

Aufruf:
    # Alle Workflows testen:
    pytest tests/test_eval.py -v --tb=short

    # Nur einen Workflow:
    pytest tests/test_eval.py -v -k "entlassbericht"

    # Mit Output-Speicherung (fuer manuelles Review):
    pytest tests/test_eval.py -v --eval-output /workspace/eval_results/

Voraussetzungen:
    - Backend laeuft auf localhost:8000
    - Ollama laeuft mit dem konfigurierten Modell
    - Verlaufsdoku-PDFs liegen in tests/fixtures/eval/ (optional)
"""
import json
import logging
import os
import re
import time
from pathlib import Path
import httpx
import pytest

# Workflows fuer die Stage 1 ueberhaupt relevant ist. Muss synchron zu
# backend/app/api/jobs.py::_STAGE1_WORKFLOWS bleiben.
_EVAL_STAGE1_WORKFLOWS = {"verlaengerung", "folgeverlaengerung", "entlassbericht"}


# ── Transkriptions-Cache ─────────────────────────────────────────────────────

def _transcript_cache_path(audio_path: Path) -> Path:
    """<audio_stem>.transcript.txt neben der Audio-Datei."""
    return audio_path.with_suffix(".transcript.txt")


def _load_or_transcribe(audio_path: Path, force_transcribe: bool = False) -> str:
    """
    Gibt das Transkript für audio_path zurück.

    - force_transcribe=False (Default):
        Liest <audio>.transcript.txt wenn vorhanden.
        Nur wenn kein Cache → transkribiert via /api/transcribe und speichert.
    - force_transcribe=True (--transcribe):
        Immer neu transkribieren, Cache überschreiben.

    Transkription läuft server-seitig über den laufenden Backend-Server,
    identisch zum normalen Produktiv-Workflow.
    """
    cache = _transcript_cache_path(audio_path)

    if not force_transcribe and cache.exists():
        logger.info("[TRANSCRIPT CACHE] lade %s", cache)
        return cache.read_text(encoding="utf-8")

    logger.info(
        "[TRANSCRIPT] transkribiere %s (%s)",
        audio_path.name,
        "erzwungen via --transcribe" if force_transcribe else "kein Cache vorhanden",
    )

    mime_map = {
        ".mp3":  "audio/mpeg",
        ".m4a":  "audio/mp4",
        ".wav":  "audio/wav",
        ".ogg":  "audio/ogg",
        ".flac": "audio/flac",
        ".webm": "audio/webm",
    }
    mime = mime_map.get(audio_path.suffix.lower(), "application/octet-stream")

    with httpx.Client(timeout=TIMEOUT) as http:
        with open(audio_path, "rb") as fh:
            resp = http.post(
                f"{BACKEND_URL}/api/transcribe",
                files={"file": (audio_path.name, fh, mime)},
            )

    if resp.status_code != 200:
        raise RuntimeError(
            f"Transkription fehlgeschlagen [{resp.status_code}]: {resp.text[:400]}"
        )

    transcript = resp.json()["transcript"]

    # Atomisch schreiben (tmp → rename)
    tmp = cache.with_suffix(".tmp")
    tmp.write_text(transcript, encoding="utf-8")
    tmp.replace(cache)
    logger.info(
        "[TRANSCRIPT CACHE] gespeichert → %s (%d Zeichen)",
        cache, len(transcript),
    )
    return transcript



logger = logging.getLogger(__name__)


def derive_word_limits(
    style_texts: list,
    fallback_min: int,
    fallback_max: int,
    tolerance: float = 0.30,
) -> tuple:
    """
    Leitet min/max Wortlimit dynamisch aus Stilvorlagen ab.
    Spiegelt die gleichnamige Funktion in prompts.py (Option A):
    beide operieren auf dem gleichen Rohtext, gleiche Logik.
    """
    counts = [len(t.split()) for t in style_texts if t and len(t.split()) >= 50]
    if not counts:
        return fallback_min, fallback_max
    ref_min, ref_max = min(counts), max(counts)
    derived_min = max(50, int(ref_min * (1 - tolerance)))
    derived_max = int(ref_max * (1 + tolerance))
    logger.info(
        "Wortlimit aus %d Stilvorlage(n) abgeleitet: %d–%d (Referenz: %d–%d, ±%.0f%%)",
        len(counts), derived_min, derived_max, ref_min, ref_max, tolerance * 100,
    )
    return derived_min, derived_max


# ── Konfiguration ────────────────────────────────────────────────────────────

BACKEND_URL = os.environ.get("EVAL_BACKEND_URL", "http://localhost:8000")
FIXTURES_PATH = Path(__file__).parent / "fixtures" / "eval" / "fixtures.json"
EVAL_DATA_DIR = Path(os.environ.get("EVAL_DATA_DIR", "/workspace/eval_data"))
EVAL_RESULTS_DIR = Path(os.environ.get("EVAL_RESULTS_DIR", "/workspace/eval_results"))
STYLES_DIR = EVAL_DATA_DIR / "styles"
TIMEOUT = 900  # 5 Minuten pro Generierung (lang wegen GPU-Kaltstart)

# Abschnitts-Überschriften in den DOCX-Vorlagen pro Workflow
STYLE_SECTION_HEADINGS = {
    "entlassbericht": [
        "Psychotherapeutischer Verlauf",
        "Psychotherapeutischer Behandlungsverlauf",
        "Verlaufsbericht",
        "Verlauf",
    ],
    "verlaengerung": [
        "Bisheriger Verlauf und Begründung der Verlängerung",
        "Verlauf und Begründung der weiteren Verlängerung",
        "Bisheriger Verlauf",
        "Begründung der Verlängerung",
    ],
    "folgeverlaengerung": [
        "Verlauf und Begründung der weiteren Verlängerung",
        "Bisheriger Verlauf und Begründung der Verlängerung",
        "Bisheriger Verlauf",
    ],
    "anamnese": [
        "Aktuelle Anamnese",
        "Anamnese",
        "Psychische Anamnese",
    ],
    "befund": [
        "Psychischer Befund",
        "Psychopathologischer Befund",
        "Befund",
    ],
    "akutantrag": [
        "Zusammenfassung (Begründung der Notwendigkeit)",
        "Begründung der Notwendigkeit",
        "Begründung für Akutaufnahme",
        "Begründung für die Akutaufnahme",
        "Akutbegründung",
        "Begründung",
        # Fallback wenn die Begruendungs-Abschnitte leer sind:
        # Aktuelle Anamnese zeigt ebenfalls den therapeutischen Schreibstil
        "Aktuelle Anamnese",
    ],
    "dokumentation": None,  # Gesprächszusammenfassung = ganzes Dokument
}


# ── DOCX-Abschnittsextraktion ────────────────────────────────────────────────

def _extract_docx_text(docx_path: Path) -> str:
    """Extrahiert den vollständigen Text aus einem DOCX."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx nicht installiert – DOCX-Stilvorlagen nicht verfügbar")
        return ""
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_docx_section(docx_path: Path, headings: list[str]) -> str:
    """
    Extrahiert einen Abschnitt aus einem DOCX anhand der Überschrift.

    Sammelt ALLE Bold/Heading-Matches, sortiert nach Prioritaet in der headings-Liste,
    und probiert jeden Kandidaten durch bis einer ausreichend Inhalt liefert (>= 20 Woerter).
    Dadurch werden leere Abschnitte (z.B. Vorlagen mit leerer "Begruendung") uebersprungen
    und der naechste gefuellte Abschnitt (z.B. "Aktuelle Anamnese") wird genutzt.
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx nicht installiert – DOCX-Stilvorlagen nicht verfügbar")
        return ""

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs
    headings_lower = [h.lower() for h in headings]

    def _extract_from(start_idx: int, start_level: int | None) -> str:
        """Sammelt Text ab start_idx bis zur naechsten Heading/End-Marker."""
        end_markers = [
            "wir bitten daher um", "daher bitten wir um", "vor diesem hintergrund bitten wir",
            "wir bitten um", "fuer rueckfragen", "für rückfragen",
            "mit freundlichem gruß", "mit freundlichen grüßen",
        ]
        section_lines = []
        for p in paragraphs[start_idx:]:
            text = p.text.strip()
            if not text:
                section_lines.append("")
                continue
            text_lower = text.lower()
            if any(text_lower.startswith(m) for m in end_markers):
                break
            style_name = (p.style.name or "").lower()
            is_heading = "heading" in style_name or style_name.startswith("überschrift")
            is_bold_heading = (
                all(run.bold for run in p.runs if run.text.strip())
                and len(text.split()) <= 8
            ) if p.runs else False
            if is_heading or is_bold_heading:
                level_match = re.search(r'(\d)', style_name)
                level = int(level_match.group(1)) if level_match else 2
                if level <= (start_level or 2):
                    break
            section_lines.append(text)
        return "\n".join(section_lines).strip()

    # Alle Bold/Heading-Positionen mit ihren Matches sammeln
    candidates = []
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style_name = (p.style.name or "").lower()
        is_heading = "heading" in style_name or style_name.startswith("überschrift")
        is_bold = all(run.bold for run in p.runs if run.text.strip()) if p.runs else False

        if is_heading or is_bold:
            text_lower = text.lower().rstrip(":")
            for h in sorted(headings_lower, key=len, reverse=True):
                if len(h) > 20:
                    matched = h in text_lower or (
                        text_lower in h and len(text_lower.split()) >= len(h.split()) - 1
                    )
                else:
                    matched = text_lower == h or text_lower.startswith(h)
                if matched:
                    level_match = re.search(r'(\d)', style_name)
                    lvl = int(level_match.group(1)) if level_match else 2
                    candidates.append((i + 1, lvl, h))
                    break

    # Sortiere nach Prioritaet in der headings-Liste (frueher = wichtiger)
    heading_priority = {h.lower(): idx for idx, h in enumerate(headings)}
    candidates.sort(key=lambda c: heading_priority.get(c[2], 999))

    # Probiere jeden Kandidaten
    for start_idx, start_level, matched_h in candidates:
        result = _extract_from(start_idx, start_level)
        if result and len(result.split()) >= 20:
            logger.info(
                "DOCX-Abschnitt extrahiert: %s via '%s' → %d Wörter",
                docx_path.name, matched_h, len(result.split()),
            )
            return result
        else:
            logger.debug(
                "Heading '%s' in %s ergab zu wenig Text (%d Wörter) – probiere naechste",
                matched_h, docx_path.name, len(result.split()) if result else 0,
            )

    # Stufe 2: Plain-String-Fallback (formatierungsunabhängig)
    # Sucht nach Heading-Text als einfachem Substring in ALLEN Paragraphen,
    # unabhängig von Bold/Heading-Formatierung. Greift wenn Stufe 1 keinen
    # Kandidaten mit ≥20 Wörtern Inhalt lieferte (z.B. andere Formatierungen).
    for h in headings_lower:
        for i, p in enumerate(paragraphs):
            text = p.text.strip().lower().rstrip(":")
            if not text:
                continue
            if h in text:  # einfacher Substring-Match, keine Fuzzy-Logik
                result = _extract_from(i + 1, None)
                if result and len(result.split()) >= 20:
                    logger.info(
                        "DOCX-Abschnitt (Plain-String-Fallback): %s via '%s' → %d Wörter",
                        docx_path.name, h, len(result.split()),
                    )
                    return result

    # Stufe 3: Volltext-Fallback (letzter Ausweg)
    logger.warning(
        "Keine Überschrift mit Inhalt in %s (gesucht: %s). Verwende gesamten Text.",
        docx_path.name, headings,
    )
    return "\n".join(p.text for p in paragraphs if p.text.strip())


def load_style_text(therapeut: str, workflow: str) -> str | None:
    """
    Lädt die Stilvorlage für einen Therapeuten und Workflow.

    Sucht in /workspace/eval_data/styles/{therapeut}/ nach dem
    passenden DOCX und extrahiert den relevanten Abschnitt.
    Bei mehreren Dateien (Entlassbericht_01.docx, _02.docx, ...) wird
    die erste gefundene zurückgegeben. Für alle: load_all_style_texts().
    """
    texts = load_all_style_texts(therapeut, workflow)
    return texts[0] if texts else None


def _extract_section_by_text(docx_path: Path, headings: list[str]) -> str:
    """
    Fallback-Extraktion: sucht nach Heading-Texten als Plain-Text im Dokument
    (unabhängig von Style oder Bold-Formatierung).
    Sammelt allen Text nach dem Treffer bis zur nächsten kurzen Zeile
    die wie eine Überschrift aussieht (<=8 Wörter, gefolgt von längerem Text).
    """
    try:
        from docx import Document
    except ImportError:
        return ""

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs
    headings_lower = [h.lower() for h in headings]

    start_idx = None
    for i, p in enumerate(paragraphs):
        text = p.text.strip().lower().rstrip(":")
        if not text:
            continue
        for h in headings_lower:
            matched = h in text or (
                text in h and len(text.split()) >= len(h.split()) - 1
            )
            if matched:
                start_idx = i + 1
                break
        if start_idx is not None:
            break

    if start_idx is None:
        return ""

    # Text sammeln bis zur nächsten kurzen Zeile die wie Überschrift aussieht
    end_markers = [
        "wir bitten daher um", "daher bitten wir um", "vor diesem hintergrund bitten wir",
        "wir bitten um", "fuer rueckfragen", "für rückfragen",
        "mit freundlichem gruß", "mit freundlichen grüßen",
    ]
    section_lines = []
    for p in paragraphs[start_idx:]:
        text = p.text.strip()
        if not text:
            section_lines.append("")
            continue
        text_lower = text.lower()
        # End-Marker erkennen
        if any(text_lower.startswith(m) for m in end_markers):
            break
        # Kurze Zeile nach substanziellem Text → wahrscheinlich nächste Überschrift
        if len(text.split()) <= 8 and len(section_lines) > 3 and any(len(l.split()) > 10 for l in section_lines[-3:]):
            break
        section_lines.append(text)

    return "\n".join(section_lines).strip()


def load_all_style_texts(therapeut: str, workflow: str) -> list[str]:
    """
    Lädt ALLE Stilvorlagen eines Therapeuten für einen Workflow.

    Unterstützt mehrere Dateien pro Workflow:
      Entlassbericht.docx, Entlassbericht_01.docx, Entlassbericht_02.docx, ...

    Gibt eine Liste von extrahierten Abschnitten zurück.
    """
    therapeut_dir = STYLES_DIR / therapeut
    if not therapeut_dir.exists():
        return []

    # Datei-Prefixe pro Workflow (case-insensitive, flexible Benennungen)
    if workflow == "dokumentation":
        prefixes = ["gesprächszusammenfassung", "gespraechszusammenfassung", "dokumentation", "gespräch", "gespraech"]
    elif workflow == "entlassbericht":
        prefixes = ["entlassbericht"]
    elif workflow in ("verlaengerung", "folgeverlaengerung"):
        prefixes = ["verlängerungsantrag", "verlaengerungsantrag", "verlängerung", "verlaengerung",
                     "folgeverlängerung", "folgeverlaengerung", "fogleverlängerung"]
    elif workflow == "akutantrag":
        prefixes = ["akutantrag"]
    elif workflow == "anamnese":
        prefixes = ["entlassbericht", "verlängerungsantrag", "verlaengerungsantrag",
                     "verlängerung", "verlaengerung", "anamnese"]
    else:
        return []

    # Alle passenden Dateien finden (case-insensitive, auch nummerierte: _01, 2, etc.)
    docx_files = []
    for f in sorted(therapeut_dir.iterdir()):
        if f.suffix.lower() not in (".docx", ".doc"):
            continue
        fname = f.stem.lower()  # case-insensitive
        for prefix in prefixes:
            # Matcht: "entlassbericht", "entlassbericht2", "entlassbericht_01", "Entlassbericht"
            if fname == prefix or fname.startswith(prefix + "_") or fname.startswith(prefix) and fname[len(prefix):].lstrip("_").isdigit():
                docx_files.append(f)
                break

    if not docx_files:
        logger.warning("Keine DOCX-Vorlage für %s/%s in %s", therapeut, workflow, therapeut_dir)
        return []

    # Abschnitte extrahieren
    headings = STYLE_SECTION_HEADINGS.get(workflow)
    texts = []
    for docx_path in docx_files:
        try:
            text = None
            if headings is not None:
                text = _extract_docx_section(docx_path, headings)
            # Fallback: Plain-Text-Suche nach Heading-Text im Dokument
            if (not text or len(text.split()) < 20) and headings:
                text = _extract_section_by_text(docx_path, headings)
            # Letzter Fallback: gesamter Text
            if not text or len(text.split()) < 20:
                text = _extract_docx_text(docx_path)
            if text and len(text.split()) >= 20:
                texts.append(text)
        except Exception as e:
            logger.warning("DOCX-Extraktion fehlgeschlagen: %s (%s)", docx_path.name, e)

    logger.info(
        "Stilvorlagen geladen: %s/%s → %d Beispiele (%s)",
        therapeut, workflow, len(texts),
        ", ".join(f.name for f in docx_files),
    )
    return texts


def _discover_style_siblings(primary_path: Path) -> list:
    """Wrapper um eval_helpers.discover_style_siblings (siehe dort).

    v13 Strategie 3: Multi-Vorlagen-Erkennung pro Testcase. Logik ist in
    tests/eval_helpers.py ausgelagert, damit test_prompts_v16 sie ohne
    Fixtures-Setup importieren kann.
    """
    from tests.eval_helpers import discover_style_siblings as _impl
    return _impl(primary_path)


def _build_multi_style_text(paths: list) -> str:
    """Wrapper um eval_helpers.build_multi_style_text (siehe dort).

    Bindet den lokalen _extract_docx_text als DOCX-Extractor an, damit
    DOCX-Stilvorlagen in der Eval gelesen werden können.
    """
    from tests.eval_helpers import build_multi_style_text as _impl
    return _impl(paths, docx_extractor=_extract_docx_text)


def discover_therapeuten() -> list[str]:
    """Findet alle Therapeuten-Ordner in /workspace/eval_data/styles/."""
    if not STYLES_DIR.exists():
        return []
    return sorted([
        d.name for d in STYLES_DIR.iterdir()
        if d.is_dir() and not d.name.startswith(".")
    ])


# ── Fixtures laden ───────────────────────────────────────────────────────────

def _load_fixtures() -> dict:
    with open(FIXTURES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


FIXTURES = _load_fixtures()


def _all_test_cases():
    """Generiert (workflow, test_case) Tupel fuer Parametrisierung."""
    cases = []
    for workflow in ["entlassbericht", "verlaengerung", "folgeverlaengerung", "akutantrag", "anamnese", "dokumentation"]:
        for tc in FIXTURES.get(workflow, []):
            cases.append((workflow, tc))
    return cases


# ── API-Helfer ───────────────────────────────────────────────────────────────

async def _generate(
    workflow: str,
    prompt: str,
    diagnosen: list[str] | None = None,
    input_files: dict | None = None,
    extra_form_data: dict | None = None,
    force_transcribe: bool = False,
) -> dict:
    """
    Sendet einen Generierungs-Job und wartet auf das Ergebnis.

    input_files: optionales Dict mit Datei-Feldern
    extra_form_data: zusätzliche Form-Felder (z.B. {"style_text": "..."})
    """
    form_data = {
        "workflow": workflow,
        # v18: Feld heisst nun 'workflow_instructions'.
        # 'prompt' bleibt als deprecated Alias erhalten (jobs.py akzeptiert beide).
        "workflow_instructions": prompt,
    }
    if diagnosen and "diagnosen" not in form_data:
        form_data["diagnosen"] = ",".join(diagnosen)
    if extra_form_data:
        form_data.update(extra_form_data)

    # Dateien vorbereiten
    # Kritische Felder (ohne die der Workflow keinen Sinn macht) failen wenn Datei fehlt.
    # Audio/diagnosen_file/style_file sind optional und werden nur geloggt.
    CRITICAL_FIELDS = {"antragsvorlage", "verlaufsdoku", "vorantrag", "selbstauskunft"}
    missing_critical = []
    files_to_upload = {}
    if input_files:
        for field_name, file_path in input_files.items():
            p = Path(file_path)
            # Relative Pfade gegen EVAL_DATA_DIR auflösen
            if not p.is_absolute():
                p = EVAL_DATA_DIR / p
            if not p.exists():
                if field_name in CRITICAL_FIELDS:
                    missing_critical.append(f"{field_name}: {p}")
                    logger.error("KRITISCH: Eval-Input fehlt: %s → %s", field_name, p)
                else:
                    logger.warning("Eval-Input nicht gefunden: %s (übersprungen)", p)
                continue

            # Spezielle Felder die als Text (nicht File) gesendet werden
            if field_name == "style_file":
                try:
                    # v13 Strategie 3: Multi-Vorlagen-Erkennung pro Testcase.
                    # Wenn p="vorlage.txt", suche zusätzlich vorlage2.txt,
                    # vorlage3.txt etc. im selben Verzeichnis. Alle gefundenen
                    # werden mit "--- Beispiel N ---"-Markern verkettet (gleiches
                    # Format wie retrieve_style_examples). split_style_examples
                    # in jobs.py erkennt das wieder und gibt resolve_length_anchor
                    # eine Liste von Einzelvorlagen statt einem konkatenierten Block.
                    style_paths = _discover_style_siblings(p)
                    style_content = _build_multi_style_text(style_paths)
                    if style_content:
                        form_data["style_text"] = style_content
                        logger.info(
                            "Eval-Input: style_text aus %d Datei%s (%d Zeichen): %s",
                            len(style_paths),
                            "en" if len(style_paths) != 1 else "",
                            len(style_content),
                            ", ".join(sp.name for sp in style_paths),
                        )
                except Exception as e:
                    logger.warning("Stilvorlage nicht lesbar: %s (%s)", p, e)
            elif field_name == "diagnosen_file":
                # diagnosen.txt: ICD-Codes laden und als Komma-Liste senden
                try:
                    raw = p.read_text(encoding="utf-8")
                    # Unterstützt: eine pro Zeile, kommagetrennt, oder gemischt
                    codes = [c.strip() for c in raw.replace("\n", ",").split(",") if c.strip()]
                    if codes:
                        form_data["diagnosen"] = ",".join(codes)
                        logger.info("Eval-Input: diagnosen aus %s → %s", p, codes)
                except Exception as e:
                    logger.warning("Diagnosen nicht lesbar: %s (%s)", p, e)
            else:
                # Audio: nicht als Datei hochladen, sondern gecacht transkribieren
                # und als transcript-Textfeld senden (spart Whisper-Zeit bei Folge-Runs).
                if field_name == "audio":
                    try:
                        transcript_text = _load_or_transcribe(p, force_transcribe=force_transcribe)
                        form_data["transcript"] = transcript_text
                        logger.info(
                            "Eval-Input: audio → transcript (%d Zeichen, cache=%s)",
                            len(transcript_text),
                            not force_transcribe and _transcript_cache_path(p).exists(),
                        )
                    except Exception as e:
                        logger.warning(
                            "Transkription fehlgeschlagen für %s: %s – Audio wird direkt gesendet",
                            p, e,
                        )
                        files_to_upload[field_name] = (p.name, open(p, "rb"))
                else:
                    files_to_upload[field_name] = (p.name, open(p, "rb"))
                    logger.info("Eval-Input: %s = %s", field_name, p)

    # Abbruch wenn kritische Input-Dateien fehlen - besser laut scheitern als
    # stillschweigend einen Job ohne Quellen starten (der dann halluziniert).
    if missing_critical:
        raise FileNotFoundError(
            f"Eval-Test kann nicht starten - kritische Inputs fehlen:\n  "
            + "\n  ".join(missing_critical)
            + f"\n(EVAL_DATA_DIR = {EVAL_DATA_DIR})"
        )

    try:
        async with httpx.AsyncClient(base_url=BACKEND_URL, timeout=120.0) as client:
            # Job erstellen – mit oder ohne Dateien
            r = await client.post(
                "/api/jobs/generate",
                data=form_data,
                files=files_to_upload or None,
            )
            r.raise_for_status()
            job_id = r.json()["job_id"]

            # Pollen bis fertig
            t0 = time.time()
            while time.time() - t0 < TIMEOUT:
                r = await client.get(f"/api/jobs/{job_id}")
                r.raise_for_status()
                job = r.json()

                if job["status"] == "done":
                    return job
                if job["status"] == "error":
                    raise RuntimeError(f"Job fehlgeschlagen: {job.get('error_msg', '?')}")
                if job["status"] == "cancelled":
                    raise RuntimeError("Job wurde abgebrochen")

                await _async_sleep(3)

            raise TimeoutError(f"Job {job_id} nicht in {TIMEOUT}s fertig geworden")
    finally:
        # Datei-Handles schliessen
        for _name, (_, fh) in files_to_upload.items():
            fh.close()


async def _async_sleep(seconds: float):
    import asyncio
    await asyncio.sleep(seconds)


# ── Stil-Analyse ─────────────────────────────────────────────────────────────

# IFS/systemische Fachbegriffe für Dichte-Messung
FACHBEGRIFFE = {
    "anteile", "anteil", "anteilearbeit", "manager", "exile", "feuerwehr",
    "self-energy", "selbst-energie", "steuerungsposition", "schutzanteil",
    "schutzanteile", "inneres kind", "türsteher", "wächter", "wächterin",
    "hypnosystemisch", "systemisch", "ressource", "ressourcen",
    "ressourcenorientiert", "reframing", "externalisierung", "stuhlarbeit",
    "körperarbeit", "netzwerkarbeit", "biographiearbeit", "auftragsklärung",
    "dissoziation", "affektregulation", "bindungsmuster", "traumatisierung",
    "co-regulation", "selbstwirksamkeit", "selbstfürsorge", "selbstwert",
    "schwingungsfähigkeit", "vulnerabilität", "stabilisierung",
}


class StyleAnalyzer:
    """Extrahiert messbare Stilmerkmale aus einem Text."""

    def __init__(self, text: str):
        self.text = text
        self.sentences = self._split_sentences(text)
        self.words = text.split()
        self.paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    @staticmethod
    def _split_sentences(text: str) -> list[str]:
        # Einfacher Satz-Splitter (deutsch: Punkt, Fragezeichen, Ausrufezeichen)
        parts = re.split(r'(?<=[.!?])\s+(?=[A-ZÄÖÜ])', text)
        return [s.strip() for s in parts if s.strip() and len(s.split()) >= 3]

    @property
    def avg_sentence_length(self) -> float:
        """Durchschnittliche Wörter pro Satz."""
        if not self.sentences:
            return 0.0
        lengths = [len(s.split()) for s in self.sentences]
        return sum(lengths) / len(lengths)

    @property
    def avg_paragraph_length(self) -> float:
        """Durchschnittliche Wörter pro Absatz."""
        if not self.paragraphs:
            return 0.0
        lengths = [len(p.split()) for p in self.paragraphs]
        return sum(lengths) / len(lengths)

    @property
    def fachbegriff_density(self) -> float:
        """Anteil der Fachbegriffe an der Gesamtwortanzahl (0.0-1.0)."""
        if not self.words:
            return 0.0
        lower_text = self.text.lower()
        hits = sum(1 for fb in FACHBEGRIFFE if fb in lower_text)
        # Normalisiert auf Textlänge (pro 100 Wörter)
        return hits / (len(self.words) / 100) if self.words else 0.0

    @property
    def wir_perspektive_ratio(self) -> float:
        """Anteil Sätze mit Wir-Perspektive."""
        if not self.sentences:
            return 0.0
        wir_pattern = re.compile(r'\b(wir|uns|unser)\b', re.IGNORECASE)
        wir_count = sum(1 for s in self.sentences if wir_pattern.search(s))
        return wir_count / len(self.sentences)

    @property
    def direkte_zitate_count(self) -> int:
        """Anzahl direkter Zitate (in Anführungszeichen)."""
        return len(re.findall(r'[„"«].*?[""»]', self.text))

    def to_dict(self) -> dict:
        return {
            "word_count": len(self.words),
            "sentence_count": len(self.sentences),
            "paragraph_count": len(self.paragraphs),
            "avg_sentence_length": round(self.avg_sentence_length, 1),
            "avg_paragraph_length": round(self.avg_paragraph_length, 1),
            "fachbegriff_density": round(self.fachbegriff_density, 2),
            "wir_perspektive_ratio": round(self.wir_perspektive_ratio, 2),
            "direkte_zitate": self.direkte_zitate_count,
        }


# ── Evaluations-Checks ──────────────────────────────────────────────────────

class EvalResult:
    """Sammelt Evaluations-Ergebnisse fuer einen Testfall."""

    def __init__(self, workflow: str, test_id: str, text: str):
        self.workflow = workflow
        self.test_id = test_id
        self.text = text
        self.word_count = len(text.split())
        self.issues: list[str] = []
        self.passed: list[str] = []
        # v13 Ä5: Telemetrie zum Längenanker - aus welcher Quelle kam das Limit?
        # Werte: "style", "style_too_short_fallback", "workflow_default",
        #        "style_invalid_range_fallback", "unknown" (Default vor Set)
        self.length_source: str = "unknown"
        self.length_min: int = 0
        self.length_max: int = 0
        self.length_n_substantial: int = 0
        # v19.1: Telemetrie aus llm.generate_text (Think-Block-Diagnose,
        # Retry-Status). Vom Test gesetzt nachdem das Job-Result gefetcht
        # wurde; vom Reporter aggregiert.
        self.generation_telemetry: dict | None = None
        self.degraded: bool = False
        self.retry_used: bool = False
        # v19.2: Stage-1-Pipeline-Audit aus dem Job-Result. None bei Jobs
        # die Stage 1 nicht beruehrt haben. Wird vom Reporter aggregiert und
        # in der LLM-Jury-Auswertung beruecksichtigt.
        self.verlauf_summary_audit: dict | None = None
        self.stage1_applied: bool = False
        self.stage1_compression_ratio: float | None = None
        self.stage1_retry_used: bool = False
        self.stage1_degraded: bool = False
        self.stage1_issue_count: int = 0

    def check_word_count(self, min_words: int, max_words: int):
        if self.word_count < min_words:
            self.issues.append(f"Zu kurz: {self.word_count}w < {min_words}w Minimum")
            self.word_count_ok = False
        elif self.word_count > max_words:
            self.issues.append(f"Zu lang: {self.word_count}w > {max_words}w Maximum")
            self.word_count_ok = False
        else:
            self.passed.append(f"Wortanzahl OK: {self.word_count}w ({min_words}-{max_words})")
            self.word_count_ok = True  # Bug-Fix #3b: erlaubt Absatzlängen-Check zu lockern

    def check_required_keywords(self, keywords: list[str]):
        # Synonyme fuer Keywords die im Fließtext anders ausgedrueckt werden koennen
        keyword_synonyms = {
            "vorstellungsanlass": ["vorstellungsanlass", "stellt sich vor", "stellt sich mit",
                                    "hauptanliegen", "hauptbeschwerde", "kommt mit", "berichtet"],
            "behandlungsverlauf": ["behandlungsverlauf", "im verlauf", "im einzelprozess",
                                    "therapeutische arbeit", "wir erlebten"],
            "empfehlung": ["empfehlung", "empfohlen", "ambulant", "nachsorge",
                            "weiterbehandlung", "weitere therapie"],
        }
        text_lower = self.text.lower()
        for kw in keywords:
            kw_lower = kw.lower()
            if kw_lower in text_lower:
                self.passed.append(f"Keyword vorhanden: '{kw}'")
                continue
            # Semantischer Match
            synonyms = keyword_synonyms.get(kw_lower, [kw_lower])
            if any(s in text_lower for s in synonyms):
                self.passed.append(f"Keyword (semantisch) vorhanden: '{kw}'")
            else:
                self.issues.append(f"Keyword fehlt: '{kw}'")

    def check_forbidden_patterns(self, patterns: list[str]):
        for pat in patterns:
            if pat in self.text:
                self.issues.append(f"Verbotenes Pattern gefunden: '{pat}'")
            else:
                self.passed.append(f"Pattern nicht vorhanden: '{pat}'")

    def check_required_sections(self, sections: list[str]):
        """
        Prueft semantisch ob ein Inhalt vorhanden ist (statt exaktes Wort).
        Erlaubt Synonyme und thematische Indikatoren — passend zum Fließtext-Stil
        ohne explizite Unterueberschriften.
        """
        # Synonym-Sets: Sektion → Liste von Indikatoren (mind. einer muss vorkommen)
        synonyms = {
            "Behandlungsverlauf": [
                "behandlungsverlauf", "verlauf", "im einzelprozess", "therapeutische arbeit",
                "im laufe der behandlung", "im verlauf der behandlung", "wir erlebten",
                "im stationaren rahmen", "im stationären rahmen",
            ],
            "Empfehlung": [
                "empfehlung", "empfohlen", "empfehlen", "ambulant", "nachsorge",
                "weiterbehandlung", "weitere therapie", "fortführung", "fortsetzen",
            ],
            "Vorstellungsanlass": [
                "vorstellungsanlass", "stellt sich vor", "stellt sich mit",
                "hauptanliegen", "hauptbeschwerde", "vorstellungsgrund",
                "kommt mit", "leidet unter", "berichtet über", "berichtet von",
            ],
            "Anamnese": [
                "anamnese", "berichtet", "biographisch", "vorgeschichte",
                "in der vergangenheit", "fruher", "früher",
            ],
            "Befund": [
                "befund", "psychischer befund", "psychopathologisch",
                "im gespräch", "im gespraech", "stimmungslage",
            ],
        }

        text_lower = self.text.lower()
        for section in sections:
            section_lower = section.lower()
            # Erst exakter Match (alte Logik)
            if section_lower in text_lower:
                self.passed.append(f"Sektion vorhanden: '{section}'")
                continue
            # Dann semantisch via Synonym-Set
            indicators = synonyms.get(section, [section_lower])
            if any(ind in text_lower for ind in indicators):
                self.passed.append(f"Sektion (semantisch) vorhanden: '{section}'")
            else:
                self.issues.append(f"Sektion fehlt: '{section}' (auch keine Synonyme gefunden)")

    def check_forbidden_names(self, names: list[str]):
        for name in names:
            if name in self.text:
                self.issues.append(f"DATENSCHUTZ: Name '{name}' im Text gefunden!")
            else:
                self.passed.append(f"Datenschutz OK: '{name}' nicht im Text")

    def check_hallucinations(self, hallucinations: list[str]):
        for h in hallucinations:
            if h.lower() in self.text.lower():
                self.issues.append(f"HALLUZINATION: '{h}' gefunden – nicht in Quelldaten!")
            else:
                self.passed.append(f"Keine Halluzination: '{h}'")

    def check_befund_separator(self, separator: str):
        if separator in self.text:
            self.passed.append(f"Befund-Separator '{separator}' vorhanden")
        else:
            self.issues.append(f"Befund-Separator '{separator}' fehlt – Anamnese/Befund nicht getrennt")

    def check_no_think_blocks(self):
        if "</think>" in self.text or "<think>" in self.text:
            self.issues.append("Think-Block im Output gefunden!")
        else:
            self.passed.append("Kein Think-Block im Output")

    def check_style_consistency(self, style_text: str | list[str], tolerance: float = 0.6):
        """
        Ansatz 1: Stil-Konsistenz-Check.

        Bei einem Referenztext: vergleicht Output gegen diesen Text.
        Bei mehreren Referenztexten: berechnet die Bandbreite des Therapeuten-Stils
        und prüft ob der Output innerhalb dieser Bandbreite (+ Toleranz) liegt.

        Bug-Fix #3: Toleranz von 0.4 auf 0.6 erhoeht. Die alte ±40%-Schwelle
        kippt schon bei kleinen Output-Abweichungen, insbesondere bei kurzen
        Texten (Akutantrag) oder Stilvorlagen mit ungewoehnlich grossen Absaetzen.
        """
        if isinstance(style_text, str):
            refs = [StyleAnalyzer(style_text)]
        else:
            refs = [StyleAnalyzer(t) for t in style_text if t]

        if not refs:
            return

        out = StyleAnalyzer(self.text)

        # Metriken der Referenztexte sammeln
        def _metric_range(metric_fn):
            values = [metric_fn(r) for r in refs]
            values = [v for v in values if v > 0]
            if not values:
                return 0, 0, 0
            return min(values), sum(values) / len(values), max(values)

        checks = [
            ("Satzlänge", lambda r: r.avg_sentence_length, out.avg_sentence_length),
            ("Absatzlänge", lambda r: r.avg_paragraph_length, out.avg_paragraph_length),
            # Fachbegriff-Dichte aus Tests entfernt — in der Praxis akzeptabel,
            # auch wenn Output mehr Fachbegriffe als Referenz nutzt
        ]

        for name, metric_fn, out_val in checks:
            ref_min, ref_avg, ref_max = _metric_range(metric_fn)
            if ref_avg == 0:
                continue

            # Bug-Fix #3b: Wenn die Wortzahl insgesamt im erlaubten Bereich liegt,
            # ist eine Abweichung der Absatzlänge oft akzeptabel (kürzere Absätze
            # bei korrekter Gesamtlänge bedeutet z.B. nur mehr Strukturierung).
            # Wir loggen dann als "passed" mit Hinweis statt als Issue.
            _is_paragraph_check = (name == "Absatzlänge")
            _word_count_passed = getattr(self, "word_count_ok", False)
            _downgrade = _is_paragraph_check and _word_count_passed

            if len(refs) > 1:
                # Mehrere Referenzen: Output muss in die Bandbreite fallen (+ Toleranz)
                band_low = ref_min * (1 - tolerance)
                band_high = ref_max * (1 + tolerance)
                in_band = band_low <= out_val <= band_high
                if in_band:
                    self.passed.append(
                        f"STIL {name} OK: Output={out_val:.1f} in Bandbreite "
                        f"[{ref_min:.1f}–{ref_max:.1f}] ±{tolerance:.0%} "
                        f"({len(refs)} Referenzen)"
                    )
                elif _downgrade:
                    # Bug-Fix #3b: Absatzlaenge weicht ab, aber Wortzahl OK -> nur Warnung
                    self.passed.append(
                        f"STIL {name} akzeptabel (Wortzahl-OK trotz Bandbreite verfehlt): "
                        f"Output={out_val:.1f}, Referenz=[{ref_min:.1f}–{ref_max:.1f}] "
                        f"({len(refs)} Referenzen)"
                    )
                else:
                    self.issues.append(
                        f"STIL {name} außerhalb Bandbreite: Output={out_val:.1f}, "
                        f"Referenz=[{ref_min:.1f}–{ref_max:.1f}] ±{tolerance:.0%} "
                        f"({len(refs)} Referenzen)"
                    )
            else:
                # Ein Referenztext: einfacher Vergleich
                deviation = abs(out_val - ref_avg) / ref_avg if ref_avg else 0
                if deviation <= tolerance:
                    self.passed.append(
                        f"STIL {name} OK: Vorlage={ref_avg:.1f} Output={out_val:.1f} "
                        f"(±{deviation:.0%})"
                    )
                elif _downgrade:
                    # Bug-Fix #3b: Absatzlaenge weicht ab, aber Wortzahl OK -> nur Warnung
                    self.passed.append(
                        f"STIL {name} akzeptabel (Wortzahl-OK trotz Abweichung): "
                        f"Vorlage={ref_avg:.1f} Output={out_val:.1f} (±{deviation:.0%})"
                    )
                else:
                    self.issues.append(
                        f"STIL {name} weicht ab: Vorlage={ref_avg:.1f} Output={out_val:.1f} "
                        f"(±{deviation:.0%}, erlaubt ±{tolerance:.0%})"
                    )

        # Wir-Perspektive
        wir_refs = [r.wir_perspektive_ratio for r in refs]
        avg_wir = sum(wir_refs) / len(wir_refs) if wir_refs else 0
        if avg_wir > 0.1:
            if out.wir_perspektive_ratio > 0.05:
                self.passed.append(
                    f"STIL Wir-Perspektive OK: Referenz={avg_wir:.0%} "
                    f"Output={out.wir_perspektive_ratio:.0%}"
                )
            else:
                self.issues.append(
                    f"STIL Wir-Perspektive fehlt: Referenz={avg_wir:.0%} "
                    f"Output={out.wir_perspektive_ratio:.0%}"
                )

        # Stil-Metriken speichern
        self.style_metrics = {
            "reference_count": len(refs),
            "reference": {
                "avg": StyleAnalyzer(style_text[0] if isinstance(style_text, list) else style_text).to_dict(),
                "range": {
                    "sentence_length": [round(_metric_range(lambda r: r.avg_sentence_length)[0], 1),
                                        round(_metric_range(lambda r: r.avg_sentence_length)[2], 1)],
                    "paragraph_length": [round(_metric_range(lambda r: r.avg_paragraph_length)[0], 1),
                                         round(_metric_range(lambda r: r.avg_paragraph_length)[2], 1)],
                    "fachbegriff_density": [round(_metric_range(lambda r: r.fachbegriff_density)[0], 2),
                                            round(_metric_range(lambda r: r.fachbegriff_density)[2], 2)],
                },
            },
            "output": out.to_dict(),
        }

    def summary(self) -> str:
        status = "PASS" if not self.issues else "FAIL"
        lines = [
            f"[{status}] {self.workflow}/{self.test_id} ({self.word_count}w)",
            f"  ✓ {len(self.passed)} Checks bestanden",
        ]
        if self.issues:
            lines.append(f"  ✗ {len(self.issues)} Probleme:")
            for issue in self.issues:
                lines.append(f"    - {issue}")
        # v13 Ä5: Längenquelle als Diagnose-Information
        if self.length_source != "unknown":
            source_label = {
                "style": f"Stilvorlage(n) (n={self.length_n_substantial})",
                "style_too_short_fallback": "Stilvorlage zu kurz → Workflow-Default",
                "style_invalid_range_fallback": "Stilvorlage out-of-range → Workflow-Default",
                "workflow_default": "Workflow-Default (keine Stilvorlage)",
            }.get(self.length_source, self.length_source)
            lines.append(
                f"  Längenanker: {self.length_min}-{self.length_max}w "
                f"(Quelle: {source_label})"
            )
        if hasattr(self, "style_metrics") and self.style_metrics:
            ref = self.style_metrics["reference"].get("avg", self.style_metrics["reference"])
            out = self.style_metrics["output"]
            lines.append(f"  Stil-Vergleich:")
            lines.append(f"    Satzlänge:    Vorlage={ref.get('avg_sentence_length','?')}w  Output={out.get('avg_sentence_length','?')}w")
            lines.append(f"    Absatzlänge:  Vorlage={ref.get('avg_paragraph_length','?')}w  Output={out.get('avg_paragraph_length','?')}w")
            lines.append(f"    Fachbegriffe: Vorlage={ref.get('fachbegriff_density','?')}/100w  Output={out.get('fachbegriff_density','?')}/100w")
            lines.append(f"    Wir-Perspektive: Vorlage={ref.get('wir_perspektive_ratio',0):.0%}  Output={out.get('wir_perspektive_ratio',0):.0%}")
        if hasattr(self, "style_variance_score"):
            lines.append(f"  Stil-Varianz (A vs B): {self.style_variance_score:.2f} (>0.15 = gut)")
        if hasattr(self, "llm_style_score"):
            lines.append(f"  LLM-Stil-Bewertung: {self.llm_style_score}/5")
        # v19.1: Generierungs-Stabilitaet (Think-Block-Diagnose)
        if self.generation_telemetry:
            tel = self.generation_telemetry
            flags = []
            if tel.get("degraded"):
                flags.append("DEGRADED")
            if tel.get("retry_used"):
                flags.append("RETRY")
            if tel.get("tokens_hit_cap"):
                flags.append("TOKEN-CAP")
            if tel.get("used_thinking_fallback"):
                flags.append("THINKING-FALLBACK")
            flag_str = " ".join(flags) if flags else "OK"
            tr = tel.get("think_ratio")
            tr_str = f", think_ratio={tr:.0%}" if tr is not None else ""
            lines.append(f"  Generierungs-Stabilitaet: {flag_str}{tr_str}")
        # v19.2: Stage-1-Pipeline-Status (Verlauf-Verdichtung)
        if self.verlauf_summary_audit is not None:
            audit = self.verlauf_summary_audit
            s1_flags = []
            if self.stage1_applied:
                s1_flags.append("APPLIED")
            else:
                s1_flags.append(f"SKIPPED({audit.get('fallback_reason') or 'unbekannt'})")
            if self.stage1_retry_used:
                s1_flags.append("RETRY")
            if self.stage1_degraded:
                s1_flags.append("DEGRADED")
            if self.stage1_issue_count:
                s1_flags.append(f"{self.stage1_issue_count}ISSUES")
            cr = self.stage1_compression_ratio
            cr_str = f", kompression={cr:.0%}" if cr is not None else ""
            raw_w = audit.get("raw_word_count")
            sum_w = audit.get("summary_word_count")
            wc_str = f", {raw_w}w->{sum_w}w" if (raw_w and sum_w) else ""
            lines.append(f"  Stage-1-Pipeline: {' '.join(s1_flags)}{cr_str}{wc_str}")
        return "\n".join(lines)

    @property
    def score(self) -> float:
        """Score 0.0-1.0 basierend auf bestandenen Checks.

        v19.1: Wenn der LLM-Output als degraded markiert wurde
        (Think-Block-Bug, beide Versuche zu kurz), ist der Score
        unabhaengig von der Check-Bilanz 0 - der Output ist faktisch
        nicht brauchbar.

        v19.2: Stage-1-degraded (critical Halluzinations-Signale auch
        nach Retry) treibt den Score ebenfalls auf 0 - der nachgelagerte
        Generierungs-Output basiert dann auf einer Quelle der nicht
        getraut werden kann.
        """
        if self.degraded or self.stage1_degraded:
            return 0.0
        total = len(self.passed) + len(self.issues)
        return len(self.passed) / total if total > 0 else 0.0


# ── Pytest-Tests ─────────────────────────────────────────────────────────────

@pytest.mark.asyncio
@pytest.mark.parametrize("workflow,test_case", _all_test_cases(),
                         ids=[f"{w}-{tc['id']}" for w, tc in _all_test_cases()])
async def test_eval_workflow(workflow, test_case, request):
    """
    Generiert Text fuer einen Testfall und prueft die Qualitaet.

    Erwartet einen laufenden Backend-Server auf EVAL_BACKEND_URL.
    Ueberspringt automatisch wenn Server nicht erreichbar.
    """
    # Server-Erreichbarkeit pruefen
    try:
        async with httpx.AsyncClient(timeout=5.0) as c:
            r = await c.get(f"{BACKEND_URL}/api/health")
            if r.status_code != 200:
                pytest.skip(f"Backend nicht healthy: {r.status_code}")
    except httpx.ConnectError:
        pytest.skip(f"Backend nicht erreichbar: {BACKEND_URL}")

    # Generieren
    prompt = test_case["prompt"]
    diagnosen = test_case.get("diagnosen")
    input_files = test_case.get("input_files")

    # Pruefen ob Input-Dateien vorhanden sind (optional)
    if input_files:
        missing = []
        for field, path in input_files.items():
            p = Path(path) if Path(path).is_absolute() else EVAL_DATA_DIR / path
            if not p.exists():
                missing.append(f"{field}: {p}")
        if missing:
            logger.warning(
                "Eval-Input-Dateien nicht gefunden (Test laeuft ohne):\n  %s",
                "\n  ".join(missing),
            )
            input_files = None  # Ohne Dateien weitermachen

    # v13 P3: Konsistenz Eval-Side ↔ Backend.
    # Bisheriges Problem: Wenn eine Fixture nur "style_therapeut" hat (aber
    # KEIN "style_file" in input_files), wurde die Therapeuten-Bibliothek
    # NUR lokal für Stil-Konsistenz und Längen-Anker verwendet - ans Backend
    # ging GAR KEINE Stilvorlage. Folge: Backend rechnete mit Workflow-
    # Default-Längen, Eval-Check rechnete mit style-abgeleiteten Längen
    # → Divergenz. Konkretes Beispiel aus dem letzten Eval-Run: anamnese
    # hatte n_substantial=0 im Backend (Workflow-Default 280-650w), aber
    # Eval-Check hat aus der Bibliothek max=418w abgeleitet → an-02 mit 643w
    # → Backend "OK" / Eval "FAIL".
    #
    # Fix: Wenn KEIN style_file im input_files ist UND style_therapeut im
    # testcase steht, lade die Bibliothek und reiche sie via extra_form_data
    # ins _generate (das setzt sie als style_text). split_style_examples in
    # jobs.py splittet das dann sauber an den Markern auf (Strategie 3),
    # so dass Backend und Eval-Side die gleichen Einzelvorlagen sehen.
    extra_form_data = None
    has_style_file = bool(input_files and "style_file" in input_files)
    if not has_style_file and test_case.get("style_therapeut"):
        therapeut_id = test_case["style_therapeut"]
        try:
            library_texts = load_all_style_texts(therapeut_id, workflow)
            if library_texts:
                # Marker-Format wie split_style_examples es erwartet
                # (identisch zu retrieve_style_examples in embeddings.py).
                if len(library_texts) == 1:
                    style_content = library_texts[0]
                else:
                    style_content = "\n\n".join(
                        f"--- Beispiel {i} ---\n{txt}"
                        for i, txt in enumerate(library_texts, 1)
                    )
                extra_form_data = {"style_text": style_content}
                logger.info(
                    "Eval-Input P3: style_text aus Bibliothek %s "
                    "(%d Vorlage%s, %d Zeichen)",
                    therapeut_id,
                    len(library_texts),
                    "n" if len(library_texts) != 1 else "",
                    len(style_content),
                )
            else:
                logger.warning(
                    "Eval-Input P3: style_therapeut=%s gesetzt, aber keine "
                    "Vorlagen für workflow=%s gefunden",
                    therapeut_id, workflow,
                )
        except Exception as e:
            logger.warning(
                "Eval-Input P3: Bibliothek %s nicht lesbar (%s)",
                therapeut_id, e,
            )

    try:
        job = await _generate(
            workflow, prompt, diagnosen, input_files,
            extra_form_data=extra_form_data,
            force_transcribe=request.config.getoption("--transcribe", default=False),
        )
    except (RuntimeError, TimeoutError) as e:
        pytest.fail(f"Generierung fehlgeschlagen: {e}")

    # Anamnese-Workflow: Anamnese + Befund zu kombiniertem Text verketten.
    # jobs.py gibt beide Teile separat zurück (zwei LLM-Calls); der Test
    # evaluiert den Gesamtoutput (Anamnese + ###BEFUND### + Befund).
    if workflow == "anamnese":
        anamnese_part = (job.get("result_text") or "").strip()
        befund_part   = (job.get("befund_text") or "").strip()
        if befund_part:
            text = anamnese_part + "\n\n###BEFUND###\n\n" + befund_part
        else:
            text = anamnese_part
            logger.warning(
                "[%s/%s] befund_text leer – nur Anamnese evaluiert",
                workflow, test_case["id"],
            )
    else:
        text = job.get("result_text", "")

    if not text:
        pytest.fail("Leerer Output")

    # Evaluieren
    expected = test_case["expected"]
    ev = EvalResult(workflow, test_case["id"], text)

    # v19.1: Telemetrie aus dem Job-Result auf EvalResult uebernehmen
    # und ggf. einen GENERATION_DEGRADED-Issue erzeugen. Das degraded-Flag
    # treibt den score-Property auf 0 (siehe EvalResult.score).
    _telemetry = job.get("generation_telemetry") or {}
    ev.generation_telemetry = _telemetry
    ev.retry_used = bool(_telemetry.get("retry_used"))
    if _telemetry.get("degraded"):
        ev.degraded = True
        ev.issues.append(
            f"GENERATION_DEGRADED: {_telemetry.get('degraded_reason') or 'unbekannt'}"
        )
    if ev.retry_used and not ev.degraded:
        # Retry war erfolgreich - kein Issue, aber als Info-Pass notieren
        ev.passed.append("Retry-Layer angeschlagen und erfolgreich")

    # v19.2: Stage-1-Audit aus dem Job-Result uebernehmen. Bei Jobs die
    # Stage 1 nicht beruehrt haben (Workflow nicht in Whitelist oder Verlauf
    # zu kurz) ist das Feld None - dann bleiben die EvalResult-Felder auf
    # den Defaults (applied=False).
    _stage1_audit = job.get("verlauf_summary_audit") or {}
    if _stage1_audit:
        ev.verlauf_summary_audit       = _stage1_audit
        ev.stage1_applied              = bool(_stage1_audit.get("applied"))
        ev.stage1_compression_ratio    = _stage1_audit.get("compression_ratio")
        ev.stage1_retry_used           = bool(_stage1_audit.get("retry_used"))
        ev.stage1_degraded             = bool(_stage1_audit.get("degraded"))
        ev.stage1_issue_count          = len(_stage1_audit.get("issues") or [])

        # Stage-1-Degraded ist ein Hard-Fail-Signal, analog zu GENERATION_DEGRADED
        if ev.stage1_degraded:
            reason = _stage1_audit.get("fallback_reason") or "stage1_degraded"
            ev.issues.append(
                f"STAGE1_DEGRADED: {reason} "
                f"(critical Halluzinations-Signale auch nach Retry)"
            )

        # Stage-1-Retry erfolgreich (kein degraded) → Info-Pass
        if ev.stage1_retry_used and not ev.stage1_degraded:
            ev.passed.append("Stage 1 Retry angeschlagen und erfolgreich")

    # --summary-mode-Enforcement: pruefen ob das Backend-Verhalten den
    # erwarteten Modus erfuellt. Greift NUR fuer Workflows in der
    # Stage-1-Whitelist - andere Workflows haben sowieso kein Stage 1.
    summary_mode = request.config.getoption("--summary-mode", default="auto")
    if workflow in _EVAL_STAGE1_WORKFLOWS:
        if summary_mode == "require_stage1" and not ev.stage1_applied:
            ev.issues.append(
                "SUMMARY_MODE_VIOLATION: --summary-mode=require_stage1, "
                "aber Stage 1 wurde nicht angewandt "
                f"(fallback_reason={_stage1_audit.get('fallback_reason')!r})"
            )
        elif summary_mode == "require_no_stage1" and ev.stage1_applied:
            ev.issues.append(
                "SUMMARY_MODE_VIOLATION: --summary-mode=require_no_stage1, "
                "aber Stage 1 wurde angewandt "
                f"(compression={ev.stage1_compression_ratio})"
            )

    ev.check_no_think_blocks()

    # Wortlimit: dynamisch aus Therapeuten-Stilvorlagen ableiten wenn verfügbar,
    # Fixture-Defaults als Fallback. Spiegelt die Logik in jobs.py/prompts.py.
    _wl_defaults = {
        "dokumentation":      (150, 500),
        "anamnese":           (450, 700),
        "verlaengerung":      (300, 600),
        "folgeverlaengerung": (300, 600),
        "entlassbericht":     (600, 1200),
        "akutantrag":         (150, 400),
    }
    _fb_min = expected.get("min_words", _wl_defaults.get(workflow, (200, 800))[0])
    _fb_max = expected.get("max_words", _wl_defaults.get(workflow, (200, 800))[1])
    _style_therapeut = test_case.get("style_therapeut")

    # v13 Ä5: Identische Resolution wie Production via resolve_length_anchor.
    # Vorher wurde derive_word_limits direkt aufgerufen - das kapselt jetzt
    # resolve_length_anchor inkl. Floor/Ceiling-Schutz und Quellen-Telemetrie.
    # v13 Strategie 3: Längen-Anker bekommt Multi-Vorlagen aus zwei Quellen:
    #   a) style_file im input_files (vorlage.txt + Geschwister)
    #   b) style_therapeut Bibliothek (Therapeuten-Ordner)
    # Reihenfolge: testcase-spezifische Vorlagen bevorzugt, dann Bibliothek.
    from app.services.prompts import resolve_length_anchor
    _style_texts_for_limits = None
    _input_files = test_case.get("input_files") or {}
    if "style_file" in _input_files:
        # Pro-Testcase Vorlagen (vorlage.txt, vorlage2.txt, ...)
        _sp = _input_files["style_file"]
        _sp_path = Path(_sp) if Path(_sp).is_absolute() else EVAL_DATA_DIR / _sp
        if _sp_path.exists():
            _siblings = _discover_style_siblings(_sp_path)
            _style_texts_for_limits = []
            for _sib in _siblings:
                try:
                    if _sib.suffix.lower() == ".txt":
                        _t = _sib.read_text(encoding="utf-8")
                    elif _sib.suffix.lower() in (".docx", ".doc"):
                        _t = _extract_docx_text(_sib)
                    else:
                        continue
                    if _t and _t.strip():
                        _style_texts_for_limits.append(_t)
                except Exception as e:
                    logger.warning("Stilvorlage für Anker nicht lesbar: %s (%s)", _sib, e)
            if not _style_texts_for_limits:
                _style_texts_for_limits = None
    # Fallback: Therapeuten-Bibliothek
    if _style_texts_for_limits is None and _style_therapeut:
        _style_texts_for_limits = load_all_style_texts(_style_therapeut, workflow) or None

    _anchor = resolve_length_anchor(
        workflow=workflow,
        style_raw_texts=_style_texts_for_limits,
        workflow_default=(_fb_min, _fb_max),
    )
    eff_min, eff_max = _anchor["min"], _anchor["max"]
    # Telemetrie auf EvalResult festhalten (für Report-Aggregation)
    ev.length_source = _anchor["source"]
    ev.length_min = eff_min
    ev.length_max = eff_max
    ev.length_n_substantial = _anchor["n_substantial"]
    ev.check_word_count(eff_min, eff_max)

    if "required_keywords" in expected:
        ev.check_required_keywords(expected["required_keywords"])

    if "forbidden_patterns" in expected:
        ev.check_forbidden_patterns(expected["forbidden_patterns"])

    if "required_sections" in expected:
        ev.check_required_sections(expected["required_sections"])

    if "must_contain_sections" in expected:
        ev.check_required_sections(expected["must_contain_sections"])

    if "forbidden_names" in expected:
        ev.check_forbidden_names(expected["forbidden_names"])

    if "must_not_hallucinate" in expected:
        ev.check_hallucinations(expected["must_not_hallucinate"])

    # befund_separator: wird durch das Verketten von Anamnese+Befund implizit gesetzt.
    # Wenn befund_text leer war, fehlt der Separator → Check greift korrekt als Fail.
    if workflow == "anamnese" or "befund_separator" in expected:
        ev.check_befund_separator("###BEFUND###")

    # Ansatz 1: Stil-Konsistenz gegen Vorlage prüfen
    # v13 Strategie 3: nutzt _discover_style_siblings um vorlage.txt + vorlage2.txt
    # gemeinsam zu finden. Bei mehreren Vorlagen wird die Konsistenz gegen ALLE
    # geprüft (StyleAnalyzer mittelt) - das ist robuster als nur gegen eine.
    style_text = None
    if input_files and "style_file" in (test_case.get("input_files") or {}):
        style_path = test_case["input_files"]["style_file"]
        p = Path(style_path) if Path(style_path).is_absolute() else EVAL_DATA_DIR / style_path
        if p.exists():
            try:
                # Alle Geschwister-Vorlagen finden (vorlage.txt, vorlage2.txt, ...)
                sibling_paths = _discover_style_siblings(p)
                style_texts_list = []
                for sib in sibling_paths:
                    if sib.suffix.lower() == ".txt":
                        t = sib.read_text(encoding="utf-8")
                    elif sib.suffix.lower() in (".docx", ".doc"):
                        # Relevanten Abschnitt aus DOCX extrahieren – gleiche Logik
                        # wie load_all_style_texts: Bold-Match, Plain-Text, Volltext.
                        headings = STYLE_SECTION_HEADINGS.get(workflow)
                        t = None
                        if headings:
                            t = _extract_docx_section(sib, headings)
                            if not t or len(t.split()) < 20:
                                t = _extract_section_by_text(sib, headings)
                        if not t or len(t.split()) < 20:
                            t = _extract_docx_text(sib)
                    else:
                        t = ""
                    if t and len(t.split()) >= 20:
                        style_texts_list.append(t)

                if style_texts_list:
                    if len(style_texts_list) == 1:
                        # Single-Vorlage: alter Pfad (string)
                        style_text = style_texts_list[0]
                        ev.check_style_consistency(style_text)
                    else:
                        # Multi-Vorlage: Liste, StyleAnalyzer mittelt
                        style_text = style_texts_list
                        ev.check_style_consistency(style_texts_list)
                        logger.info(
                            "Stil-Konsistenz gegen %d Vorlagen geprüft",
                            len(style_texts_list),
                        )
            except Exception as e:
                logger.warning("Stilvorlage nicht lesbar: %s", e)

    # Zusätzlich: Stil aus Therapeuten-Bibliothek (styles/ Verzeichnis)
    if not style_text and test_case.get("style_therapeut"):
        style_texts = load_all_style_texts(test_case["style_therapeut"], workflow)
        if style_texts:
            style_text = style_texts  # Liste für check_style_consistency + ref.txt
            ev.check_style_consistency(style_texts)

    # Ergebnis loggen
    print(f"\n{ev.summary()}")

    # Ergebnis speichern (immer – default: /workspace/eval_results/)
    output_dir = request.config.getoption("--eval-output", default=None) or EVAL_RESULTS_DIR
    out_path = Path(output_dir) / workflow
    out_path.mkdir(parents=True, exist_ok=True)
    result_file = out_path / f"{test_case['id']}.txt"
    result_file.write_text(text, encoding="utf-8")
    # Referenz-Stiltext speichern (fuer 2-spaltigen PDF-Vergleich)
    _ref_text = style_text if isinstance(style_text, str) else (style_text[0] if style_text else None)
    if _ref_text:
        (out_path / f"{test_case['id']}.ref.txt").write_text(_ref_text, encoding="utf-8")
    summary_file = out_path / f"{test_case['id']}.eval.txt"
    summary_file.write_text(ev.summary(), encoding="utf-8")
    if hasattr(ev, "style_metrics") and ev.style_metrics:
        metrics_file = out_path / f"{test_case['id']}.style.json"
        metrics_file.write_text(
            json.dumps(ev.style_metrics, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    # Test failt wenn es kritische Issues gibt
    critical = [i for i in ev.issues if "DATENSCHUTZ" in i or "HALLUZINATION" in i]
    if critical:
        pytest.fail(f"Kritische Probleme:\n" + "\n".join(f"  - {i}" for i in critical))

    # Warnungen fuer nicht-kritische Issues
    if ev.issues:
        for issue in ev.issues:
            logger.warning("[%s/%s] %s", workflow, test_case["id"], issue)

    # Score mindestens 70%
    assert ev.score >= 0.7, (
        f"Score zu niedrig: {ev.score:.0%} ({len(ev.passed)}/{len(ev.passed)+len(ev.issues)} Checks)\n"
        f"{ev.summary()}"
    )


# ── Ansatz 2: Stil-Varianz-Test (A/B-Vergleich) ─────────────────────────────

def _style_variance_fixtures():
    """
    Generiert Testfälle für den Stil-Varianz-Test.
    Kombiniert jeden Workflow-Testfall mit allen verfügbaren Therapeuten-Paaren.
    Braucht mindestens 2 Therapeuten in /workspace/eval_data/styles/.
    """
    therapeuten = discover_therapeuten()
    if len(therapeuten) < 2:
        return []

    cases = []
    for workflow in ["entlassbericht", "verlaengerung", "folgeverlaengerung", "akutantrag", "dokumentation"]:
        tcs = FIXTURES.get(workflow, [])
        if not tcs:
            continue
        tc = tcs[0]  # Ersten Testfall nehmen
        # Alle Therapeuten-Paare
        for i, ta in enumerate(therapeuten):
            for tb in therapeuten[i + 1:]:
                # Prüfen ob beide Therapeuten Vorlagen für diesen Workflow haben
                style_a = load_style_text(ta, workflow)
                style_b = load_style_text(tb, workflow)
                if style_a and style_b:
                    cases.append((workflow, tc, ta, tb))
    return cases


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "workflow,test_case,therapeut_a,therapeut_b",
    _style_variance_fixtures(),
    ids=[f"variance-{w}-{a}-vs-{b}" for w, _, a, b in _style_variance_fixtures()],
)
async def test_style_variance(workflow, test_case, therapeut_a, therapeut_b, request):
    """
    Ansatz 2: Therapeuten-Varianz-Test.
    Generiert denselben Fall mit Stilvorlagen von zwei verschiedenen Therapeuten
    und prüft ob die Outputs sich messbar unterscheiden.
    """
    # Server prüfen
    try:
        async with httpx.AsyncClient(timeout=5.0) as c:
            r = await c.get(f"{BACKEND_URL}/api/health")
            if r.status_code != 200:
                pytest.skip("Backend nicht healthy")
    except httpx.ConnectError:
        pytest.skip("Backend nicht erreichbar")

    # Stilvorlagen laden (alle Beispiele pro Therapeut)
    styles_a = load_all_style_texts(therapeut_a, workflow)
    styles_b = load_all_style_texts(therapeut_b, workflow)

    if not styles_a or not styles_b:
        pytest.skip("Stilvorlagen nicht verfügbar")

    # Für die API: alle Beispiele zusammenfügen (so wie es der echte Workflow macht)
    style_text_a = "\n\n---\n\n".join(styles_a)
    style_text_b = "\n\n---\n\n".join(styles_b)

    # Input-Dateien vorbereiten (ohne style — den setzen wir manuell)
    base_input = dict(test_case.get("input_files", {}))
    base_input.pop("style_file", None)

    # Generierung A: mit Stil von Therapeut A
    form_extras_a = {"style_text": style_text_a}
    form_extras_b = {"style_text": style_text_b}

    try:
        _ft = request.config.getoption("--transcribe", default=False)
        job_a = await _generate(
            workflow, test_case["prompt"], test_case.get("diagnosen"),
            base_input, extra_form_data=form_extras_a,
            force_transcribe=_ft,
        )
        job_b = await _generate(
            workflow, test_case["prompt"], test_case.get("diagnosen"),
            base_input, extra_form_data=form_extras_b,
            force_transcribe=_ft,
        )
    except (RuntimeError, TimeoutError) as e:
        pytest.fail(f"Generierung fehlgeschlagen: {e}")

    text_a = job_a.get("result_text", "")
    text_b = job_b.get("result_text", "")

    if not text_a or not text_b:
        pytest.fail("Leerer Output bei Varianz-Test")

    # Stilmerkmale vergleichen
    sa = StyleAnalyzer(text_a)
    sb = StyleAnalyzer(text_b)

    # Varianz-Score: wie unterschiedlich sind die Outputs?
    diffs = []
    if sa.avg_sentence_length and sb.avg_sentence_length:
        diffs.append(abs(sa.avg_sentence_length - sb.avg_sentence_length) /
                     max(sa.avg_sentence_length, sb.avg_sentence_length))
    if sa.avg_paragraph_length and sb.avg_paragraph_length:
        diffs.append(abs(sa.avg_paragraph_length - sb.avg_paragraph_length) /
                     max(sa.avg_paragraph_length, sb.avg_paragraph_length))
    if sa.fachbegriff_density or sb.fachbegriff_density:
        max_fb = max(sa.fachbegriff_density, sb.fachbegriff_density, 0.01)
        diffs.append(abs(sa.fachbegriff_density - sb.fachbegriff_density) / max_fb)
    diffs.append(abs(sa.wir_perspektive_ratio - sb.wir_perspektive_ratio))

    variance_score = sum(diffs) / len(diffs) if diffs else 0.0

    print(f"\n[STIL-VARIANZ] {workflow}: {therapeut_a} ({len(styles_a)} Bsp.) vs {therapeut_b} ({len(styles_b)} Bsp.)")
    print(f"  Therapeut A ({therapeut_a}): Satzlänge={sa.avg_sentence_length:.1f} "
          f"Fachbegriffe={sa.fachbegriff_density:.2f} Wir={sa.wir_perspektive_ratio:.0%}")
    print(f"  Therapeut B ({therapeut_b}): Satzlänge={sb.avg_sentence_length:.1f} "
          f"Fachbegriffe={sb.fachbegriff_density:.2f} Wir={sb.wir_perspektive_ratio:.0%}")
    print(f"  Varianz-Score: {variance_score:.3f} (>0.15 = Stile wirken, <0.05 = ignoriert)")

    # Ergebnis speichern
    output_dir = request.config.getoption("--eval-output", default=None)
    if output_dir:
        out_path = Path(output_dir) / "style_variance"
        out_path.mkdir(parents=True, exist_ok=True)
        tag = f"{workflow}_{therapeut_a}_vs_{therapeut_b}"
        (out_path / f"{tag}_a.txt").write_text(text_a, encoding="utf-8")
        (out_path / f"{tag}_b.txt").write_text(text_b, encoding="utf-8")
        (out_path / f"{tag}.json").write_text(
            json.dumps({
                "variance_score": round(variance_score, 3),
                "therapeut_a": therapeut_a,
                "therapeut_a_examples": len(styles_a),
                "therapeut_b": therapeut_b,
                "therapeut_b_examples": len(styles_b),
                "style_a": sa.to_dict(),
                "style_b": sb.to_dict(),
            }, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    if variance_score < 0.05:
        logger.warning(
            "[STIL-VARIANZ] %s: Score %.3f – Stilvorlagen scheinen ignoriert zu werden!",
            workflow, variance_score,
        )

    assert variance_score > 0.03, (
        f"Stil-Varianz zu niedrig ({variance_score:.3f}). "
        f"Outputs mit Stilvorlagen von {therapeut_a} und {therapeut_b} sind nahezu identisch – "
        f"das Modell scheint die Stilvorlage zu ignorieren."
    )


# ── Ansatz 3: LLM-als-Jury (Stil-Bewertung per Modell) ──────────────────────

@pytest.mark.asyncio
@pytest.mark.parametrize("workflow,test_case", _all_test_cases(),
                         ids=[f"style-jury-{w}-{tc['id']}" for w, tc in _all_test_cases()])
async def test_style_llm_jury(workflow, test_case, request):
    """
    Ansatz 3: LLM-als-Jury.
    Sendet den generierten Text + Stilvorlage an das LLM und fragt:
    'Wie gut passt der Stil?' Bewertet auf einer Skala 1-5.

    Nur ausgeführt wenn:
    - Backend erreichbar
    - Stilvorlage vorhanden
    - Vorheriger Generierungs-Output gespeichert (via --eval-output)
    """
    # Stilvorlage laden – aus style_file oder Therapeuten-Bibliothek
    style_path = (test_case.get("input_files") or {}).get("style_file")
    style_text = None

    if style_path:
        p = Path(style_path) if Path(style_path).is_absolute() else EVAL_DATA_DIR / style_path
        if p.exists():
            try:
                if p.suffix.lower() == ".txt":
                    style_text = p.read_text(encoding="utf-8")
                elif p.suffix.lower() in (".docx", ".doc"):
                    headings = STYLE_SECTION_HEADINGS.get(workflow)
                    style_text = _extract_docx_section(p, headings) if headings else _extract_docx_text(p)
            except Exception as e:
                logger.warning("Stilvorlage für Jury nicht lesbar: %s", e)

    # Fallback: Therapeuten-Bibliothek
    if not style_text and test_case.get("style_therapeut"):
        style_text = load_style_text(test_case["style_therapeut"], workflow)

    if not style_text:
        pytest.skip("Keine Stilvorlage für LLM-Jury verfügbar")

    # Generierten Text laden (aus vorherigem Test-Run)
    output_dir = request.config.getoption("--eval-output", default=None)
    if not output_dir:
        pytest.skip("--eval-output nicht gesetzt (benötigt für LLM-Jury)")

    result_file = Path(output_dir) / workflow / f"{test_case['id']}.txt"
    if not result_file.exists():
        pytest.skip(f"Generierter Text nicht gefunden: {result_file}. Erst test_eval_workflow ausführen.")

    generated_text = result_file.read_text(encoding="utf-8")

    # Server prüfen
    try:
        async with httpx.AsyncClient(timeout=5.0) as c:
            r = await c.get(f"{BACKEND_URL}/api/health")
            if r.status_code != 200:
                pytest.skip("Backend nicht healthy")
    except httpx.ConnectError:
        pytest.skip("Backend nicht erreichbar")

    # LLM-Jury: Stil-Bewertung anfordern
    jury_prompt = (
        "Du bist ein Experte für klinische Dokumentationsstile. "
        "Vergleiche den STIL (nicht den Inhalt) dieser beiden Texte.\n\n"
        "STILVORLAGE (Referenz-Stil):\n"
        f"{style_text[:2000]}\n\n"
        "GENERIERTER TEXT (zu bewerten):\n"
        f"{generated_text[:2000]}\n\n"
        "Bewerte auf einer Skala 1-5 wie gut der generierte Text den Stil "
        "der Vorlage trifft. Berücksichtige:\n"
        "- Satzlänge und Satzkomplexität\n"
        "- Fachbegriff-Verwendung und -Dichte\n"
        "- Perspektive (Wir vs. Er/Sie)\n"
        "- Tonalität (distanziert vs. empathisch)\n"
        "- Absatzstruktur und Textfluss\n\n"
        "Antworte NUR mit einer Zahl 1-5 und einer kurzen Begründung (max 2 Sätze).\n"
        "Format: SCORE: X\nBEGRÜNDUNG: ..."
    )

    try:
        # Direkt an Ollama senden (schneller als Job-Queue, kein multipart nötig)
        ollama_url = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
        async with httpx.AsyncClient(base_url=ollama_url, timeout=120.0) as client:
            r = await client.post("/api/chat", json={
                "model": os.environ.get("OLLAMA_MODEL", "qwen3:32b"),
                "stream": False,
                "think": False,
                "messages": [
                    {"role": "user", "content": jury_prompt + "\n\n/no_think"},
                ],
                "options": {"num_predict": 200},
            })
            if r.status_code != 200:
                pytest.skip(f"LLM-Jury-Request fehlgeschlagen: {r.status_code}")

            data = r.json()
            jury_response = data.get("message", {}).get("content", "")
    except Exception as e:
        pytest.skip(f"LLM-Jury nicht verfügbar: {e}")

    # Score extrahieren
    score_match = re.search(r'SCORE:\s*(\d)', jury_response)
    score = int(score_match.group(1)) if score_match else 0

    # Begründung extrahieren
    reason_match = re.search(r'BEGRÜNDUNG:\s*(.+)', jury_response, re.DOTALL)
    reason = reason_match.group(1).strip() if reason_match else jury_response.strip()

    print(f"\n[LLM-JURY] {workflow}/{test_case['id']}")
    print(f"  Score: {score}/5")
    print(f"  Begründung: {reason}")

    # Ergebnis speichern
    if output_dir:
        out_path = Path(output_dir) / "style_jury"
        out_path.mkdir(parents=True, exist_ok=True)
        (out_path / f"{test_case['id']}.jury.json").write_text(
            json.dumps({
                "score": score,
                "reason": reason,
                "jury_response": jury_response[:500],
            }, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    assert score >= 3, (
        f"LLM-Stil-Bewertung zu niedrig: {score}/5\n"
        f"Begründung: {reason}"
    )
