"""
Dokumenten-Extraktions-Service – sysTelios KI-Dokumentation
============================================================

Dreistufige Fallback-Kette fuer PDF-Extraktion:

  Stufe 1 - pdfplumber      Maschinenlesbares PDF (digital ausgefuellt)
  Stufe 2 - Tesseract OCR   Guter Scan, klare Handschrift, gedruckter Text
  Stufe 3 - Ollama Vision   Schlechter Scan, unleserliche Handschrift,
                            komplexe Formularlayouts

Qualitaetspruefung nach jeder Stufe:
  - Mindestlaenge (Zeichenanzahl)
  - Lesbarkeitsquotient (Anteil alphanumerischer Zeichen)
  - Sprachplausibilitaet (deutsche Stoppwoerter erkennbar)
  - Tesseract Confidence-Score (falls verfuegbar)
  - Wiederholungsrate (OCR-Artefakt-Erkennung)

Weitere Formate:
  DOCX  - python-docx (inkl. Tabellen, Kopf-/Fusszeilen)
  TXT   - direkte UTF-8/Latin-1-Dekodierung
  Bild  - Tesseract dann Ollama Vision Fallback

Datenschutz:
  Alle Verarbeitungsschritte laufen lokal. Kein externer API-Aufruf.
"""

from __future__ import annotations

import asyncio
import base64
import io
import logging
import re
import unicodedata
from pathlib import Path
from typing import NamedTuple

import httpx
from PIL import Image, ImageEnhance, ImageFilter

from app.core.config import settings

logger = logging.getLogger(__name__)


# ── Relevante Abschnitte pro Workflow (fuer Stilvorlagen-Extraktion) ──────────

STYLE_SECTION_HEADINGS = {
    "entlassbericht": [
        "Psychotherapeutischer Verlauf",
        "Psychotherapeutischer Behandlungsverlauf",
        "Therapie und Verlauf",
        "Verlaufsbericht",
        "Verlauf",
    ],
    "verlaengerung": [
        "Bisheriger Verlauf und Begründung der Verlängerung",
        "Verlauf und Begründung der weiteren Verlängerung",
        "Bisheriger Verlauf",
        "Begründung der Verlängerung",
    ],
    "folgeverlaengerung": [
        "Verlauf und Begründung der weiteren Verlängerung",
        "Bisheriger Verlauf und Begründung der Verlängerung",
        "Bisheriger Verlauf",
    ],
    "anamnese": [
        "Aktuelle Anamnese",
        "Anamnese",
        "Psychische Anamnese",
    ],
    "dokumentation": [
        "Auftragsklärung",
        "Relevante Gesprächsinhalte",
    ],
    "akutantrag": [
        "Zusammenfassung (Begründung der Notwendigkeit)",
        "Begründung der Notwendigkeit",
        "Begründung für Akutaufnahme",
        "Begründung für die Akutaufnahme",
        "Akutbegründung",
        "Begründung",
        # Fallback wenn die Begruendungs-Abschnitte leer sind:
        # Aktuelle Anamnese zeigt ebenfalls den therapeutischen Schreibstil
        "Aktuelle Anamnese",
    ],
}


def extract_docx_section(file_path: Path, workflow: str) -> str:
    """
    Extrahiert den relevanten Abschnitt aus einem DOCX basierend auf dem Workflow.

    Erkennt Ueberschriften sowohl als Heading-Style als auch als fettgedruckten Text.
    Bei langen Heading-Suchstrings (>20 Zeichen) wird Substring-Match erlaubt,
    bei kurzen nur exakt/startswith (vermeidet Falsch-Positive).

    Gibt den gesamten Text zurueck wenn kein relevanter Abschnitt gefunden wird.
    """
    import re

    headings = STYLE_SECTION_HEADINGS.get(workflow)
    if not headings:
        return _extract_docx_fulltext(file_path)

    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx nicht installiert")
        return ""

    doc = Document(str(file_path))
    paragraphs = doc.paragraphs
    headings_lower = [h.lower() for h in headings]

    def _extract_from(start_idx: int, start_level: int | None) -> str:
        """Sammelt Text ab start_idx bis zur naechsten Heading/End-Marker."""
        end_markers = [
            "wir bitten daher um", "daher bitten wir um", "vor diesem hintergrund bitten wir",
            "wir bitten um", "fuer rueckfragen", "für rückfragen",
            "mit freundlichem gruß", "mit freundlichen grüßen",
        ]
        section_lines = []
        for p in paragraphs[start_idx:]:
            text = p.text.strip()
            if not text:
                section_lines.append("")
                continue
            text_lower = text.lower()
            if any(text_lower.startswith(m) for m in end_markers):
                break
            style_name = (p.style.name or "").lower()
            is_heading = "heading" in style_name or style_name.startswith("überschrift")
            is_bold_heading = (
                all(run.bold for run in p.runs if run.text.strip())
                and len(text.split()) <= 8
            ) if p.runs else False
            if is_heading or is_bold_heading:
                level_match = re.search(r"(\d)", style_name)
                level = int(level_match.group(1)) if level_match else 2
                if level <= (start_level or 2):
                    break
            section_lines.append(text)
        return "\n".join(section_lines).strip()

    # Alle Bold/Heading-Positionen mit ihren Matches sammeln
    candidates = []  # [(start_idx, start_level, matched_heading), ...]
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style_name = (p.style.name or "").lower()
        is_heading = "heading" in style_name or style_name.startswith("überschrift")
        is_bold = all(run.bold for run in p.runs if run.text.strip()) if p.runs else False

        if is_heading or is_bold:
            text_lower = text.lower().rstrip(":")
            for h in sorted(headings_lower, key=len, reverse=True):
                if len(h) > 20:
                    matched = h in text_lower or (
                        text_lower in h and len(text_lower.split()) >= len(h.split()) - 1
                    )
                else:
                    matched = text_lower == h or text_lower.startswith(h)
                if matched:
                    level_match = re.search(r"(\d)", style_name)
                    lvl = int(level_match.group(1)) if level_match else 2
                    candidates.append((i + 1, lvl, h))
                    break

    # Sortiere Kandidaten nach Heading-Prioritaet (Position in headings-Liste)
    # Je frueher in der Liste, desto hoeher die Prioritaet
    heading_priority = {h.lower(): idx for idx, h in enumerate(headings)}
    candidates.sort(key=lambda c: heading_priority.get(c[2], 999))

    # Probiere jeden Kandidaten – nimm den ersten mit ausreichend Inhalt
    for start_idx, start_level, matched_h in candidates:
        result = _extract_from(start_idx, start_level)
        if result and len(result.split()) >= 20:
            logger.info("DOCX-Abschnitt: %s/%s via '%s' → %d Woerter",
                         file_path.name, workflow, matched_h, len(result.split()))
            return result
        else:
            logger.debug("Heading '%s' in %s ergab zu wenig Text (%d Woerter) – probiere naechste",
                          matched_h, file_path.name, len(result.split()) if result else 0)

    # Kein Bold-Match mit Inhalt → Plain-Text-Suche versuchen
    logger.info("Kein Bold-Match mit Inhalt in %s fuer %s – versuche Plain-Text-Suche",
                 file_path.name, workflow)
    result = _extract_section_by_text(file_path, headings)
    if result and len(result.split()) >= 20:
        logger.info("DOCX-Abschnitt via Plain-Text: %s/%s → %d Woerter",
                     file_path.name, workflow, len(result.split()))
        return result
    logger.info("Auch Plain-Text-Suche fehlgeschlagen – verwende gesamten Text")
    return _extract_docx_fulltext(file_path)


def _extract_section_by_text(file_path: Path, headings: list[str]) -> str:
    """
    Fallback-Extraktion: sucht nach Heading-Texten als Plain-Text im Dokument
    (unabhaengig von Style oder Bold-Formatierung).
    Sammelt allen Text nach dem Treffer bis zur naechsten kurzen Zeile
    die wie eine Ueberschrift aussieht (<=8 Woerter, gefolgt von laengerem Text).
    """
    try:
        from docx import Document
    except ImportError:
        return ""

    doc = Document(str(file_path))
    paragraphs = doc.paragraphs
    headings_lower = [h.lower() for h in headings]

    start_idx = None
    for i, p in enumerate(paragraphs):
        text = p.text.strip().lower().rstrip(":")
        if not text:
            continue
        # Spezifischste zuerst (laengere Headings)
        for h in sorted(headings_lower, key=len, reverse=True):
            if len(h) > 20:
                matched = h in text or (
                    text in h and len(text.split()) >= len(h.split()) - 1
                )
                if matched:
                    start_idx = i + 1
                    break
            else:
                if text == h or text.startswith(h):
                    start_idx = i + 1
                    break
        if start_idx is not None:
            break

    if start_idx is None:
        return ""

    # Text sammeln bis zur naechsten kurzen Zeile die wie Ueberschrift aussieht
    end_markers = [
        "wir bitten daher um", "daher bitten wir um", "vor diesem hintergrund bitten wir",
        "wir bitten um", "fuer rueckfragen", "für rückfragen",
        "mit freundlichem gruß", "mit freundlichen grüßen",
    ]
    section_lines = []
    for p in paragraphs[start_idx:]:
        text = p.text.strip()
        if not text:
            section_lines.append("")
            continue
        text_lower = text.lower()
        # End-Marker erkennen
        if any(text_lower.startswith(m) for m in end_markers):
            break
        # Kurze Zeile nach substanziellem Text → wahrscheinlich naechste Ueberschrift
        if len(text.split()) <= 8 and len(section_lines) > 3 and any(len(l.split()) > 10 for l in section_lines[-3:]):
            break
        section_lines.append(text)

    return "\n".join(section_lines).strip()


def extract_patient_name(text: str) -> dict | None:
    """
    Extrahiert den Patientennamen aus einem Dokument-Text.

    Sucht nach typischen Mustern im Briefkopf:
      - "Wir berichten über Herrn/Frau [Vorname] [Nachname]"
      - "Herrn/Frau\n[Vorname] [Nachname]" (mit Zeilenumbruch)
      - Selbstauskunft: "Name:" / "Nachname:" / "Vorname:"

    Rueckgabe:
      {"anrede": "Herr"/"Frau", "vorname": "...", "nachname": "...", "initial": "K."}
      oder None wenn nichts gefunden.

    Die Initiale ist der erste Buchstabe des Nachnamens + Punkt.
    """
    import re
    if not text or len(text) < 20:
        return None

    # Suche nur in den ersten 2000 Zeichen (Briefkopf/Deckblatt)
    head = text[:2000]

    # Muster 1: "wir berichten über Herrn/Frau Vorname Nachname"
    m = re.search(
        r"[Ww]ir berichten (?:über|ueber)\s+(Herrn|Frau)\s+([A-ZÄÖÜ][a-zäöüß\-]+)"
        r"(?:\s+(?:von\s+)?(?:der|de[nr])?\s*([A-ZÄÖÜ][\wäöüÄÖÜß\.\-]*(?:\s+[A-ZÄÖÜ][\wäöüÄÖÜß\.\-]*)?))",
        head,
    )
    if m:
        anrede = m.group(1)
        if anrede == "Herrn":
            anrede = "Herr"
        vorname = m.group(2)
        nachname = (m.group(3) or "").strip()
        if nachname:
            initial = nachname.split()[0][0].upper() + "."
            # Falls van/von: Initiale vom letzten Teil
            if nachname.split()[0].lower() in ("van", "von", "de", "der", "zu"):
                if len(nachname.split()) > 1:
                    initial = nachname.split()[-1][0].upper() + "."
            return {
                "anrede": anrede,
                "vorname": vorname,
                "nachname": nachname,
                "initial": initial,
            }

    # Muster 2: "Herr/Frau" auf einer Zeile, Name auf naechster Zeile (Briefkopf-Block)
    m = re.search(
        r"^\s*(Herr|Frau)\s*\n+([A-ZÄÖÜ][a-zäöüß\-]+)\s+([A-ZÄÖÜ][\wäöüÄÖÜß\.\-]+)",
        head, re.MULTILINE,
    )
    if m:
        anrede = m.group(1)
        vorname = m.group(2)
        nachname = m.group(3)
        initial = nachname[0].upper() + "."
        return {"anrede": anrede, "vorname": vorname, "nachname": nachname, "initial": initial}

    # Muster 3: Selbstauskunft "Name: ..." bzw. "Nachname: ..."
    nachname = None
    vorname = None
    anrede = None
    m_nn = re.search(r"Nachname[:\s]+([A-ZÄÖÜ][\wäöüÄÖÜß\-]+)", head)
    m_vn = re.search(r"Vorname[:\s]+([A-ZÄÖÜ][\wäöüÄÖÜß\-]+)", head)
    m_ga = re.search(r"Geschlecht[:\s]+(m[aä]nnlich|weiblich|m|w)", head, re.IGNORECASE)
    if m_nn:
        nachname = m_nn.group(1)
    if m_vn:
        vorname = m_vn.group(1)
    if m_ga:
        g = m_ga.group(1).lower()
        anrede = "Frau" if g.startswith("w") else "Herr"
    if nachname:
        if not anrede:
            # Fallback: aus Vorname schliessen (sehr grobe Heuristik)
            anrede = "Frau" if vorname and vorname.endswith("a") else "Herr"
        initial = nachname[0].upper() + "."
        return {
            "anrede": anrede,
            "vorname": vorname or "",
            "nachname": nachname,
            "initial": initial,
        }

    return None


# Generische Bezeichnungen die NIE als Patientenname akzeptiert werden duerfen.
# Wenn das Frontend so etwas als "patientenname" schickt, ist das ein Default-Wert
# oder ein versehentlich kopierter Hinweistext - kein echter Name.
# Vergleich case-insensitive nach Strip von Slash/Whitespace.
_GENERIC_NAME_BLACKLIST = frozenset({
    "klient", "klientin", "klientinnen", "klienten",
    "patient", "patientin", "patientinnen", "patienten",
    "der klient", "die klientin", "die patientin", "der patient",
    "der/die klient/in", "der/die patient/in",
    "name", "nachname", "vorname",
    "unbekannt", "anonym", "n.n.", "nn",
    "x", "y", "z",
    "frau x", "herr x", "frau y", "herr y",
    "frau x.", "herr x.", "frau y.", "herr y.",
})

# Maximalwerte fuer plausible Namen. Echte deutsche Nachnamen liegen weit darunter.
_MAX_NACHNAME_LEN = 30
_MAX_VORNAME_LEN  = 50


def parse_explicit_patient_name(name_str: str) -> dict | None:
    """
    Parst einen explizit uebergebenen Namen-String.

    Akzeptierte Formate:
      - "Herr/Frau Vorname Nachname"  → {anrede, vorname, nachname, initial}
      - "Herr/Frau Nachname"          → {anrede, nachname, initial, vorname=""}
      - "Vorname Nachname"            → {anrede="", vorname, nachname, initial}
      - "Nachname"                    → {anrede="", vorname="", nachname, initial}

    Wenn Anrede fehlt, wird sie leer gelassen (Modell nutzt dann 'die Klientin' etc).

    Lehnt generische Bezeichnungen wie "die Klientin/der Klient", "Patient" usw.
    explizit ab — solche Strings sind kein gueltiger Name und wuerden in der
    nachgelagerten Replace-Logik (build_system_prompt) als '[Patient/in]'-Ersatz
    in den finalen Prompt landen.
    """
    import re
    if not name_str:
        return None
    s = name_str.strip()
    if not s:
        return None

    # Frueher Blacklist-Check: ganzer String (case-insensitive, normalisiert)
    s_norm = re.sub(r"\s+", " ", s.replace("/", " ").replace(",", " ")).strip().lower()
    if s_norm in _GENERIC_NAME_BLACKLIST:
        return None

    # Heuristik: Jedes Vorkommen von "klient" oder "patient" im Input ist ein
    # starker Hinweis auf einen Hinweistext statt einem Eigennamen - echte
    # deutsche Nachnamen enthalten diese Substrings praktisch nie.
    s_low = s.lower()
    if "klient" in s_low or "patient" in s_low:
        return None

    anrede = ""
    # Anrede am Anfang erkennen (Herr/Frau/Herrn)
    m = re.match(r"^(Herrn?|Frau)\s+(.+)$", s)
    if m:
        anrede = m.group(1)
        if anrede == "Herrn":
            anrede = "Herr"
        rest = m.group(2).strip()
    else:
        rest = s

    parts = rest.split()
    if not parts:
        return None

    if len(parts) == 1:
        # Nur Nachname
        nachname = parts[0]
        vorname = ""
    else:
        # Annahme: letztes Wort = Nachname, Rest = Vorname(n)
        # Ausnahme: "van Dyk", "von Beethoven" etc. → letztes Wort bleibt Nachname
        nachname = parts[-1]
        vorname = " ".join(parts[:-1])

    # Sanity-Checks: zu lange "Namen" sind Hinweistexte, keine Eigennamen
    if len(nachname) > _MAX_NACHNAME_LEN:
        return None
    if vorname and len(vorname) > _MAX_VORNAME_LEN:
        return None

    # Letzter Blacklist-Check auf den extrahierten Nachnamen alleine
    # (z.B. "die Klient" -> nachname="Klient")
    if nachname.lower().rstrip(".") in _GENERIC_NAME_BLACKLIST:
        return None

    initial = nachname[0].upper() + "." if nachname else ""
    if not initial:
        return None

    return {
        "anrede": anrede,
        "vorname": vorname,
        "nachname": nachname,
        "initial": initial,
    }


def _extract_docx_fulltext(file_path: Path) -> str:
    """Extrahiert den gesamten Text eines DOCX als Fallback."""
    try:
        from docx import Document
        doc = Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.warning("DOCX-Volltext-Extraktion fehlgeschlagen: %s", e)
        return ""


MIN_CHARS = 80
MIN_READABILITY = 0.72
MIN_CONFIDENCE = 45.0

DE_STOPWORDS = frozenset([
    "und", "der", "die", "das", "ist", "ein", "eine", "zu", "in", "von",
    "mit", "auf", "nicht", "sich", "dem", "des", "den", "bei", "durch",
    "nach", "wie", "auch", "an", "oder", "hat", "wird", "sind", "kann",
    "als", "im", "am", "um", "so", "aber", "noch", "aus", "er", "sie",
    "wir", "ich", "mein", "sein", "haben", "werden", "dass", "patient",
    "klient", "therapie", "behandlung", "datum", "bericht", "name",
    "alter", "diagnose", "beschwerden", "anamnese", "befund",
])

TESS_CONFIG = "--oem 3 --psm 6 -l deu"


class ExtractionResult(NamedTuple):
    text: str
    method: str
    quality: float
    pages: int
    warnings: list


class TextQuality(NamedTuple):
    ok: bool
    score: float
    reason: str


# ── Public API ────────────────────────────────────────────────────────────────

async def extract_text(file_path: Path) -> str:
    result = await extract_text_with_meta(file_path)
    for w in result.warnings:
        logger.warning("[OCR] %s - %s", file_path.name, w)
    logger.info(
        "[OCR] %s -> Methode: %s | Qualitaet: %.0f%% | Seiten: %d",
        file_path.name, result.method, result.quality * 100, result.pages,
    )
    return result.text


async def extract_text_with_meta(file_path: Path) -> ExtractionResult:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return await _extract_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return await _extract_docx(file_path)
    elif suffix == ".txt":
        return await _extract_txt(file_path)
    elif suffix in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"):
        return await _extract_image(file_path)
    else:
        raise ValueError(
            f"Nicht unterstuetztes Dateiformat: '{suffix}'. "
            "Unterstuetzt: PDF, DOCX, TXT, JPG, PNG, TIFF, BMP, WEBP"
        )


async def extract_style_context(file_path: Path, llm_generate_fn, workflow: str = "") -> str:
    # Fuer DOCX: relevanten Abschnitt extrahieren statt ganzen Text
    if file_path.suffix.lower() in (".docx", ".doc") and workflow:
        raw_text = extract_docx_section(file_path, workflow)
    else:
        raw_text = await extract_text(file_path)
    if len(raw_text.strip()) < 100:
        return ""
    sample = raw_text[:3000]
    style_prompt = (
        "Analysiere den Schreibstil des folgenden klinischen Textes eines Therapeuten. "
        "Beschreibe in 3-5 praegnanten Saetzen: Satzlaenge, Fachbegriff-Dichte, "
        "Formulierungsgewohnheiten, Tonalitaet und besondere sprachliche Merkmale. "
        "Schreibe die Beschreibung als direkte Anweisung fuer einen anderen Schreiber "
        "(beginne mit Schreibe in einem Stil der ...)."
    )
    result = await llm_generate_fn(
        system_prompt=style_prompt,
        user_content=f"TEXT ZUM ANALYSIEREN:\n{sample}",
        max_tokens=300,
    )
    return result.get("text", "")


# ── Qualitaetspruefung ────────────────────────────────────────────────────────

def _assess_quality(text: str, min_chars: int = MIN_CHARS) -> TextQuality:
    stripped = text.strip()
    if len(stripped) < min_chars:
        return TextQuality(False, 0.0, f"Zu kurz ({len(stripped)} < {min_chars} Zeichen)")
    readable = sum(
        1 for c in stripped
        if c.isalnum() or c.isspace() or c in ".,;:!?-()"
    )
    readability = readable / len(stripped)
    if readability < MIN_READABILITY:
        return TextQuality(False, readability,
            f"Zu viele Sonderzeichen (Lesbarkeit {readability:.0%})")
    if _has_excessive_repetition(stripped):
        return TextQuality(False, 0.2, "Excessive Zeichenwiederholungen (OCR-Artefakt)")
    words = re.findall(r"\b\w+\b", stripped.lower())
    if words:
        hits = sum(1 for w in words if w in DE_STOPWORDS)
        lang_score = min(1.0, hits / max(1, len(words) * 0.05))
    else:
        lang_score = 0.0
    score = (readability * 0.5) + (lang_score * 0.3) + (min(1.0, len(stripped) / 500) * 0.2)
    if lang_score < 0.1 and len(words) > 20:
        return TextQuality(False, score,
            f"Text erscheint nicht deutsch (lang_score: {lang_score:.2f})")
    return TextQuality(True, min(1.0, score), "OK")


def _has_excessive_repetition(text: str) -> bool:
    return bool(re.search(r"([^\s\-_])\1{6,}", text))


def _normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = re.sub(r"[^\S\n\t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[\x0c\x0b]", "\n\n", text)
    return text.strip()


# ── Bild-Vorverarbeitung ──────────────────────────────────────────────────────

def _preprocess_image_for_ocr(img: Image.Image) -> Image.Image:
    img = img.convert("L")
    w, h = img.size
    if w < 1200:
        scale = max(1.5, 1800 / w)
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    img = img.filter(ImageFilter.SHARPEN)
    img = img.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
    return ImageEnhance.Contrast(img).enhance(1.8)


def _image_to_base64(img: Image.Image, max_size: int = 1600) -> str:
    w, h = img.size
    if max(w, h) > max_size:
        ratio = max_size / max(w, h)
        img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG", quality=88)
    return base64.b64encode(buf.getvalue()).decode()


# ── Stufe 1: pdfplumber ───────────────────────────────────────────────────────

def _pdfplumber_extract(file_path: Path):
    import pdfplumber
    parts = []
    with pdfplumber.open(str(file_path)) as pdf:
        pages = len(pdf.pages)
        for page in pdf.pages:
            t = page.extract_text(x_tolerance=3, y_tolerance=3)
            if t and t.strip():
                parts.append(t.strip())
            for table in page.extract_tables():
                rows = [" | ".join(str(c).strip() for c in row if c and str(c).strip())
                        for row in table]
                rows = [r for r in rows if r]
                if rows:
                    parts.append("\n".join(rows))
    return _normalize_text("\n\n".join(parts)), pages


# ── Stufe 2: Tesseract ────────────────────────────────────────────────────────

def _tesseract_extract_pdf(file_path: Path):
    import pdf2image, pytesseract
    images = pdf2image.convert_from_path(str(file_path), dpi=300, fmt="PNG")
    parts, confidences = [], []
    for img in images:
        processed = _preprocess_image_for_ocr(img)
        try:
            data = pytesseract.image_to_data(
                processed, lang="deu", config=TESS_CONFIG,
                output_type=pytesseract.Output.DICT)
            valid = [c for c in data["conf"] if isinstance(c, (int, float)) and c > 0]
            if valid:
                confidences.append(sum(valid) / len(valid))
        except Exception:
            pass
        text = pytesseract.image_to_string(processed, lang="deu", config=TESS_CONFIG)
        if text.strip():
            parts.append(text.strip())
    avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
    return _normalize_text("\n\n".join(parts)), len(images), avg_conf


def _tesseract_extract_image(img: Image.Image):
    import pytesseract
    processed = _preprocess_image_for_ocr(img)
    try:
        data = pytesseract.image_to_data(
            processed, lang="deu", config=TESS_CONFIG,
            output_type=pytesseract.Output.DICT)
        valid = [c for c in data["conf"] if isinstance(c, (int, float)) and c > 0]
        confidence = sum(valid) / len(valid) if valid else 0.0
    except Exception:
        confidence = 0.0
    text = pytesseract.image_to_string(processed, lang="deu", config=TESS_CONFIG)
    return _normalize_text(text), confidence


# ── Stufe 3: Ollama Vision ────────────────────────────────────────────────────

async def _check_vision_model_available() -> bool:
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            r = await client.get(f"{settings.OLLAMA_HOST}/api/tags")
            if r.status_code == 200:
                models = [m["name"] for m in r.json().get("models", [])]
                return any(settings.VISION_MODEL in m for m in models)
    except Exception:
        pass
    return False


async def _ollama_vision_page(b64_image: str, page_num: int, total_pages: int) -> str:
    prompt = (
        f"Dies ist Seite {page_num} von {total_pages} eines medizinischen Dokuments "
        "der sysTelios Klinik. "
        "Extrahiere den GESAMTEN Text vollstaendig und exakt. "
        "Behalte die Struktur bei. "
        "Angekreuzte Checkboxen mit [X], nicht angekreuzte mit [ ]. "
        "Nur der extrahierte Text, keine Kommentare."
    )
    payload = {
        "model": settings.VISION_MODEL,
        "prompt": prompt,
        "images": [b64_image],
        "stream": False,
        "options": {"temperature": 0.1, "num_predict": 2048},
    }
    async with httpx.AsyncClient(timeout=120.0) as client:
        r = await client.post(f"{settings.OLLAMA_HOST}/api/generate", json=payload)
        r.raise_for_status()
    return r.json().get("response", "").strip()


async def _ollama_vision_extract_pdf(file_path: Path):
    import pdf2image
    loop = asyncio.get_event_loop()
    images = await loop.run_in_executor(
        None, lambda: pdf2image.convert_from_path(str(file_path), dpi=200, fmt="PNG"))
    total = len(images)
    parts = []
    for i, img in enumerate(images, 1):
        b64 = _image_to_base64(img)
        try:
            page_text = await _ollama_vision_page(b64, i, total)
            if page_text:
                parts.append(f"[Seite {i}]\n{page_text}")
        except Exception as e:
            logger.warning("[Vision] Seite %d fehlgeschlagen: %s", i, e)
            parts.append(f"[Seite {i} - Fehler]")
    return _normalize_text("\n\n".join(parts)), total


async def _ollama_vision_extract_image(img: Image.Image) -> str:
    return await _ollama_vision_page(_image_to_base64(img), 1, 1)


# ── PDF Fallback-Kette ────────────────────────────────────────────────────────

async def _extract_pdf(file_path: Path) -> ExtractionResult:
    warnings = []
    loop = asyncio.get_event_loop()

    # Stufe 1: pdfplumber
    logger.debug("[OCR] PDF Stufe 1: pdfplumber - %s", file_path.name)
    try:
        text, pages = await loop.run_in_executor(None, _pdfplumber_extract, file_path)
        q = _assess_quality(text)
        if q.ok:
            return ExtractionResult(text, "pdfplumber", q.score, pages, warnings)
        warnings.append(f"Stufe 1 unzureichend: {q.reason}")
    except ImportError:
        warnings.append("pdfplumber nicht installiert")
    except Exception as e:
        warnings.append(f"Stufe 1 Fehler: {e}")
        logger.warning("[OCR] Stufe 1 fehlgeschlagen: %s", e)

    # Stufe 2: Tesseract
    logger.debug("[OCR] PDF Stufe 2: Tesseract - %s", file_path.name)
    try:
        text, pages, conf = await loop.run_in_executor(None, _tesseract_extract_pdf, file_path)
        q = _assess_quality(text)
        conf_ok = conf == 0 or conf >= MIN_CONFIDENCE
        if conf > 0 and not conf_ok:
            warnings.append(f"Tesseract-Konfidenz niedrig: {conf:.1f}")
        if q.ok and conf_ok:
            return ExtractionResult(text, "tesseract", q.score, pages, warnings)
        reason = q.reason if not q.ok else f"Konfidenz zu niedrig ({conf:.1f})"
        warnings.append(f"Stufe 2 unzureichend: {reason}")
    except ImportError:
        warnings.append("pytesseract/pdf2image nicht installiert")
    except Exception as e:
        warnings.append(f"Stufe 2 Fehler: {e}")
        logger.warning("[OCR] Stufe 2 fehlgeschlagen: %s", e)

    # Stufe 3: Ollama Vision
    logger.debug("[OCR] PDF Stufe 3: Ollama Vision - %s", file_path.name)
    if not await _check_vision_model_available():
        warnings.append(f"Ollama Vision nicht verfuegbar (ollama pull {settings.VISION_MODEL})")
        raise RuntimeError(
            f"PDF '{file_path.name}': Alle Extraktionsstufen fehlgeschlagen. "
            f"Details: {chr(59).join(warnings)}"
        )
    try:
        text, pages = await _ollama_vision_extract_pdf(file_path)
        q = _assess_quality(text, min_chars=50)
        if q.ok:
            return ExtractionResult(text, "ollama_vision", q.score, pages, warnings)
        warnings.append(f"Stufe 3 unzureichend: {q.reason}")
    except Exception as e:
        warnings.append(f"Stufe 3 Fehler: {e}")
        logger.error("[OCR] Stufe 3 fehlgeschlagen: %s", e)

    raise RuntimeError(
        f"PDF '{file_path.name}': Alle Extraktionsstufen fehlgeschlagen. "
        f"Details: {chr(59).join(warnings)}"
    )


# ── DOCX ─────────────────────────────────────────────────────────────────────

async def _extract_docx(file_path: Path) -> ExtractionResult:
    loop = asyncio.get_event_loop()

    def _run():
        from docx import Document
        doc = Document(str(file_path))
        parts, warns = [], []
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(f"## {t}" if para.style.name.startswith("Heading") else t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        for section in doc.sections:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
        imgs = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
        if imgs:
            warns.append(f"DOCX enthaelt {imgs} eingebettete Bilder (nicht extrahiert)")
        return _normalize_text("\n".join(parts)), warns

    try:
        text, warns = await loop.run_in_executor(None, _run)
        q = _assess_quality(text)
        return ExtractionResult(text, "docx", q.score, 1, warns)
    except Exception as e:
        raise RuntimeError(f"DOCX-Extraktion fehlgeschlagen: {e}") from e


# ── TXT ──────────────────────────────────────────────────────────────────────

async def _extract_txt(file_path: Path) -> ExtractionResult:
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        text = file_path.read_text(encoding="latin-1", errors="replace")
    text = _normalize_text(text)
    q = _assess_quality(text)
    return ExtractionResult(text, "txt", q.score, 1, [])


# ── Bild ─────────────────────────────────────────────────────────────────────

async def _extract_image(file_path: Path) -> ExtractionResult:
    warnings = []
    loop = asyncio.get_event_loop()
    try:
        img = await loop.run_in_executor(None, Image.open, str(file_path))
    except Exception as e:
        raise RuntimeError(f"Bild konnte nicht geoeffnet werden: {e}") from e

    # Stufe 1: Tesseract
    try:
        text, conf = await loop.run_in_executor(None, _tesseract_extract_image, img)
        q = _assess_quality(text)
        if q.ok and (conf == 0 or conf >= MIN_CONFIDENCE):
            return ExtractionResult(text, "image_tess", q.score, 1, warnings)
        reason = q.reason if not q.ok else f"Konfidenz {conf:.1f}"
        warnings.append(f"Tesseract unzureichend: {reason}")
    except Exception as e:
        warnings.append(f"Tesseract Fehler: {e}")

    # Stufe 2: Ollama Vision
    if not await _check_vision_model_available():
        raise RuntimeError(
            f"Bild-OCR fehlgeschlagen und Ollama Vision nicht verfuegbar. "
            f"Bitte: ollama pull {settings.VISION_MODEL}"
        )
    try:
        text = await _ollama_vision_extract_image(img)
        q = _assess_quality(text, min_chars=30)
        if q.ok:
            return ExtractionResult(text, "image_vision", q.score, 1, warnings)
        warnings.append(f"Ollama Vision unzureichend: {q.reason}")
    except Exception as e:
        warnings.append(f"Ollama Vision Fehler: {e}")

    raise RuntimeError(
        f"Bild '{file_path.name}' konnte nicht extrahiert werden. "
        f"Details: {chr(59).join(warnings)}"
    )
